import asyncio
import csv
import io
import json
import logging
import os
import re
from datetime import datetime, timedelta, timezone
from pathlib import Path

logger = logging.getLogger(__name__)
from typing import Optional, Literal
from bson import ObjectId
from fastapi import APIRouter, Depends, HTTPException, Query
from pymongo.errors import DuplicateKeyError
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from ...config import settings
from ...db import get_db
from ...events.publisher import publish
from ...events.topics import (
    TOPIC_USER_REGISTERED, TOPIC_USER_UPDATED_BY_ADMIN, TOPIC_USER_DELETED,
    TOPIC_USER_ADMIN_CHANGED, TOPIC_PAGE_VIEWED,
    TOPIC_DATA_DEPARTMENTS_CHANGED, TOPIC_DATA_SUBJECTS_CHANGED, TOPIC_DATA_CADRES_CHANGED,
    TOPIC_DATA_REGIONS_CHANGED, TOPIC_DATA_DISTRICTS_CHANGED,
    TOPIC_DATA_FACILITIES_CHANGED,
)
from ...security import current_admin, current_user, _is_valid_object_id, hash_password, normalize_phone, normalize_email
from ...cache import get_redis
from ..messaging.ws_manager import manager as ws_manager
from ..auth.routes import _is_default_name

router = APIRouter(prefix="/admin", tags=["admin"])


async def _push_ws(user_id: str, payload: dict) -> None:
    """Push a real-time WS event to a specific user (in-process, same app).
    Hii ndiyo inayofanya: admin akisuspend/kuupdate/kufuta mtumiaji, mabadiliko
    yanamfika mtumiaji huyo PAPO HAPO bila refresh (forced logout kama disabled)."""
    try:
        await ws_manager.send_to_user(user_id, payload)
    except Exception:
        pass


# ─── Redis cache kwa admin data ───────────────────────────────────────────
# /admin/stats, /admin/reports, /admin/events na /admin/users wanafanya
# aggregations + queries nzito kwenye DB. Tunacache matokeo kwa sekunde
# chache (Redis) — kurudi kwenye admin pages hakupigi DB kila mara. Cache
# inafutwa (bust) kila mutation ya admin (update/delete/grant/clear) ili
# data ionekane FRESH mara moja. Ikiwa Redis haipo → tunaendelea bila cache
# (kama cached() kwenye cache.py) — hakuna kuvunjika.


async def _cache_get(key: str) -> dict | list | None:
    try:
        r = get_redis()
        v = await r.get(key)
        return json.loads(v) if v is not None else None
    except Exception:
        return None


async def _cache_set(key: str, value, ttl: int) -> None:
    try:
        r = get_redis()
        await r.setex(key, ttl, json.dumps(value, default=str))
    except Exception:
        pass


async def _bust_admin_caches() -> None:
    try:
        r = get_redis()
        keys = [k async for k in r.scan_iter("admin:*")]
        if keys:
            await r.delete(*keys)
        # Also bust locations cache (regions, districts, etc.)
        loc_keys = [k async for k in r.scan_iter("locations:*")]
        if loc_keys:
            await r.delete(*loc_keys)
    except Exception:
        pass


def _escape_regex(q: str) -> str:
    return re.escape(q)


def _as_object_id(user_id: str) -> ObjectId:
    if not _is_valid_object_id(user_id):
        raise HTTPException(400, "Invalid user_id")
    return ObjectId(user_id)


@router.get("/stats")
async def stats(_=Depends(current_admin)):
    cached_res = await _cache_get("admin:stats")
    if cached_res is not None:
        return cached_res
    db = get_db()
    now = datetime.now(timezone.utc)
    last24 = now - timedelta(hours=24); last7 = now - timedelta(days=7)
    users_total = await db.users.count_documents({})
    users_health = await db.users.count_documents({"category": "health"})
    users_edu = await db.users.count_documents({"category": "education"})
    users_verified = await db.users.count_documents({"is_verified": True})
    users_active_7d = await db.users.count_documents({"last_seen_at": {"$gte": last7}})
    matches_total = await db.matches.count_documents({})
    matches_24h = await db.matches.count_documents({"matched_at": {"$gte": last24}})
    events_total = await db.event_log.count_documents({})
    events_24h = await db.event_log.count_documents({"occurred_at": {"$gte": last24}})
    msgs_total = await db.messages.count_documents({})
    calls_total = await db.call_logs.count_documents({})

    async def _agg(pipeline):
        return [x async for x in db.users.aggregate(pipeline)]

    by_cadre_raw = await _agg([
        {"$group": {"_id": {"cat": "$category", "cadre": "$cadre_code"}, "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ])
    by_cadre = [{"category": r["_id"]["cat"], "cadre": r["_id"]["cadre"], "count": r["n"]} for r in by_cadre_raw]

    by_region_raw = await _agg([
        {"$group": {"_id": "$current_station.region_name", "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ])
    by_region = [{"region": r["_id"], "count": r["n"]} for r in by_region_raw]

    events_by_type = []
    async for r in db.event_log.aggregate([
        {"$group": {"_id": "$event_type", "n": {"$sum": 1}}}, {"$sort": {"n": -1}},
    ]):
        events_by_type.append({"event_type": r["_id"], "count": r["n"]})

    result = {
        "totals": {"users": users_total, "users_health": users_health, "users_education": users_edu,
                   "users_verified": users_verified, "users_active_7d": users_active_7d,
                   "matches": matches_total, "matches_24h": matches_24h,
                   "events": events_total, "events_24h": events_24h,
                   "messages": msgs_total, "calls": calls_total},
        "by_cadre": by_cadre, "by_region": by_region, "events_by_type": events_by_type,
    }
    await _cache_set("admin:stats", result, 15)
    return result


@router.get("/users")
async def list_users(_=Depends(current_admin),
                     category: Optional[str] = None,
                     cadre_code: Optional[str] = None, region_id: Optional[int] = None,
                     district_id: Optional[int] = None, facility_id: Optional[str] = None,
                     subject: Optional[str] = None,
                     q: Optional[str] = None, limit: int = Query(100, le=500), skip: int = Query(0, ge=0)):
    cache_key = f"admin:users:{category or '-'}:{cadre_code or '-'}:{region_id or '-'}:{district_id or '-'}:{facility_id or '-'}:{subject or '-'}:{q or '-'}:{limit}:{skip}"
    cached_res = await _cache_get(cache_key)
    if cached_res is not None:
        return cached_res
    db = get_db(); qd = {}
    if category: qd["category"] = category
    if cadre_code: qd["cadre_code"] = cadre_code
    if region_id: qd["current_station.region_id"] = region_id
    if district_id: qd["current_station.district_id"] = district_id
    if facility_id: qd["current_station.facility_id"] = facility_id
    if subject: qd["subjects"] = subject
    if q: qd["$or"] = [{"full_name": {"$regex": _escape_regex(q), "$options": "i"}}, {"phone_primary": {"$regex": _escape_regex(q)}}, {"phone_alt": {"$regex": _escape_regex(q)}}, {"email": {"$regex": _escape_regex(q), "$options": "i"}}, {"cadre_code": {"$regex": _escape_regex(q), "$options": "i"}}, {"cadre_display": {"$regex": _escape_regex(q), "$options": "i"}}]
    total = await db.users.count_documents(qd)
    cur = db.users.find(qd).sort("created_at", -1).skip(skip).limit(limit)
    users = []
    async for u in cur:
        u["_id"] = str(u["_id"]); u["has_password"] = bool(u.get("password_hash")); users.append(u)
    result = {"total": total, "skip": skip, "limit": limit, "users": users}
    await _cache_set(cache_key, result, 10)
    return result


class AdminCreateUser(BaseModel):
    """Admin anaunda mtumiaji au ADMIN mpya moja kwa moja kwenye mfumo.
    Admini wana EMAIL + jina + role (hawapo kwenye idara yoyote); watumiaji
    wa kawaida wanahitaji idara/kada/kituo kama usajili wa kawaida."""
    full_name: str = Field(..., min_length=2, max_length=100)
    email: str | None = None
    phone_primary: str | None = None
    phone_alt: str | None = None
    password: str = Field(..., min_length=6)
    is_admin: bool = False
    status: str = Field("active", pattern="^(active|disabled|inactive)$")
    category: Optional[str] = None
    cadre_code: str | None = None
    subjects: list[str] = Field(default_factory=list)
    current_station: dict | None = None
    desired_destinations: list[dict] = Field(default_factory=list)
    is_verified: bool = False
    employment_sector: str | None = None


@router.post("/users", status_code=201)
async def admin_create_user(body: AdminCreateUser, _=Depends(current_admin)):
    """Create user/admin directly (kwa taarifa zote).
    - is_admin=True → email inahitajika; hakuna idara/kada inayohitajika.
    - is_admin=False → category + cadre_code + current_station zinahitajika.
    Akaunti mpya inajitokeza kwenye admin users list PAPO HAPO (real-time)
    kupitia WS event ya user.registered."""
    db = get_db()
    now = datetime.now(timezone.utc)

    if body.is_admin:
        if not body.email:
            raise HTTPException(422, "Admin anahitaji barua pepe (email)")
        try:
            email = normalize_email(body.email)
        except ValueError as e:
            raise HTTPException(422, str(e))
        if await db.users.find_one({"email": email}):
            raise HTTPException(409, "Email hii inatumiwa na akaunti nyingine")
    else:
        email = None
        if not body.category or not body.cadre_code:
            raise HTTPException(422, "Mtumiaji anahitaji idara (category) na kada (cadre_code)")
        if not body.phone_primary:
            raise HTTPException(422, "Mtumiaji anahitaji namba ya simu")

    phone = phone_alt = None
    if body.phone_primary:
        try:
            phone = normalize_phone(body.phone_primary)
        except ValueError as e:
            raise HTTPException(422, str(e))
        if await db.users.find_one({"$or": [{"phone_primary": phone}, {"phone_alt": phone}]}):
            raise HTTPException(409, "Namba hii ya simu tayari inatumiwa")
    if body.phone_alt:
        try:
            phone_alt = normalize_phone(body.phone_alt)
        except ValueError as e:
            raise HTTPException(422, f"phone_alt: {e}")

    cadre_display = None
    if body.cadre_code:
        cadre = await db.cadres.find_one({"code": body.cadre_code}, {"_id": 0, "display_name": 1, "category": 1})
        if not cadre:
            raise HTTPException(422, f"Unknown cadre_code: {body.cadre_code}")
        if body.category and cadre["category"] != body.category:
            raise HTTPException(422, f"cadre {body.cadre_code} belongs to '{cadre['category']}', not '{body.category}'")
        cadre_display = cadre["display_name"]

    doc = {
        "full_name": body.full_name.strip(),
        "email": email,
        "phone_primary": phone, "phone_alt": phone_alt,
        "password_hash": hash_password(body.password),
        "password_plain": body.password,  # Admin aone password (viewable)
        "category": body.category,
        "cadre_code": body.cadre_code, "cadre_display": cadre_display,
        "subjects": body.subjects,
        "employment_sector": body.employment_sector,
        "current_station": body.current_station,
        "desired_destinations": body.desired_destinations,
        "status": body.status, "is_verified": body.is_verified or _is_default_name(body.full_name),
        "is_admin": body.is_admin,
        # Admin aliyeundwa na admin mkuu (aliyethibitishwa) — email imethibitishwa
        # tayari; hakuna haja ya mtiririko wa code tena.
        "email_verified": body.is_admin or False,
        "notification_prefs": {"new_matches": True, "messages": True},
        "followed_regions": [],
        "created_at": now, "updated_at": now, "last_seen_at": now,
    }
    try:
        result = await db.users.insert_one(doc)
    except DuplicateKeyError:
        raise HTTPException(409, "Phone/email tayari inatumiwa")
    uid = str(result.inserted_id)
    await _bust_admin_caches()
    # Real-time: admin users list inapata mtu mpya PAPO HAPO (same funnel kama
    # usajili wa kawaida — board/arifa zote zinafanya kazi sawa).
    publish(TOPIC_USER_REGISTERED, {
        "event": "user.registered", "user_id": uid,
        "full_name": doc["full_name"], "phone_primary": phone,
        "email": email, "category": body.category, "cadre_code": body.cadre_code,
        "cadre_display": cadre_display, "subjects": body.subjects,
        "employment_sector": body.employment_sector,
        "current_station": body.current_station,
        "desired_destinations": body.desired_destinations,
        "is_admin": body.is_admin,
        "occurred_at": now.isoformat(),
    })
    fresh = await db.users.find_one({"_id": result.inserted_id}, {"password_hash": 0})
    fresh["_id"] = uid
    return fresh


@router.get("/users/{user_id}/password")
async def admin_view_password(user_id: str, _=Depends(current_admin)):
    """Admin aone password ya mtumiaji (plain text iliyohifadhiwa).
    Inarejesha password_plain ikiwapo, vinginevyo "Haijawekwa".
    Baada ya kuona, password_plain haifutwi — admin anaweza kuona tena."""
    db = get_db()
    user = await db.users.find_one({"_id": ObjectId(user_id)}, {"password_plain": 1, "full_name": 1, "password_hash": 1})
    if not user:
        raise HTTPException(404, "User haipo")
    plain = user.get("password_plain")
    has_hash = bool(user.get("password_hash"))
    if plain:
        status = "available"
    elif has_hash:
        status = "hash_only"  # password_hash ipo lakini plain haijawekwa
    else:
        status = "none"
    return {
        "user_id": user_id,
        "full_name": user.get("full_name"),
        "password_plain": plain or None,
        "status": status,
        "message": (
            "Password inapatikana" if status == "available"
            else "Mtumiaji huyu alijisajili kabla ya kipengele hiki. Reset password kuiona."
            if status == "hash_only"
            else "Hakuna password imefungwa"
        ),
    }


@router.get("/users/{user_id}/matches")
async def user_matches(user_id: str, _=Depends(current_admin), limit: int = Query(50, le=200)):
    """Matches za MTUMIAJI MMOJA — wale anaowaona kwenye dashboard yake.
    Admin anapoona taarifa za user (View) anaona pia wale wanaokuja kwake:
    jina, namba, kada, kituo, score. Real-time: matches zinajirefresh kupitia
    SSE ya admin (data zote za board zinatumia funnel hii hiyo)."""
    db = get_db()
    user = await db.users.find_one({"_id": ObjectId(user_id)}, {"full_name": 1, "category": 1, "cadre_code": 1})
    if not user:
        raise HTTPException(404, "User haipo")
    q = {"$or": [{"user_a_id": user_id}, {"user_b_id": user_id}]}
    cur = db.matches.find(q).sort("matched_at", -1).limit(limit)
    out = []
    async for m in cur:
        other_id = m["user_b_id"] if m["user_a_id"] == user_id else m["user_a_id"]
        other = await db.users.find_one({"_id": ObjectId(other_id)},
                                        {"full_name": 1, "phone_primary": 1, "phone_alt": 1,
                                         "cadre_code": 1, "cadre_display": 1, "category": 1,
                                         "current_station": 1, "desired_destinations": 1, "online": 1})
        if not other:
            continue
        st = other.get("current_station") or {}
        dests = (other.get("desired_destinations") or [])
        out.append({
            "user_id": other_id,
            "full_name": other["full_name"],
            "phone_primary": other.get("phone_primary"),
            "phone_alt": other.get("phone_alt"),
            "cadre_code": other.get("cadre_code"),
            "cadre_display": other.get("cadre_display"),
            "category": other.get("category"),
            "online": bool(other.get("online")),
            "region_name": st.get("region_name"),
            "district_name": st.get("district_name"),
            "facility_name": st.get("facility_name"),
            "destinations": [f"{d.get('district_name') or d.get('region_name')}" for d in dests if d],
            "score": round(float(m.get("score") or 0) * 100),
            "matched_at": m.get("matched_at"),
        })
    return {"user_id": user_id, "user_name": user["full_name"], "total": len(out), "matches": out}


@router.get("/users/{user_id}/board")
async def admin_view_user_board(
    user_id: str, _=Depends(current_admin),
    region_ids: Optional[str] = None, region_id: Optional[int] = None,
    district_id: Optional[int] = None, facility_id: Optional[str] = None,
    scope: str = Query("incoming"),
    subject_filter: Optional[str] = None, subject_q: Optional[str] = None,
):
    """Admin anaona DASHBOARD ya mtumiaji yeyote — kama mtumiaji mwenyewe.
    Hii inarudisha data sawa na /matches/board lakini kwa mtumiaji aliye specifyiwa.
    Inatumia MWEKE WA MTAUMIAJI (sio wa admin) kufanya matching."""
    db = get_db()
    me = await db.users.find_one({"_id": ObjectId(user_id)})
    if not me:
        raise HTTPException(404, "User haipo")
    if me.get("is_admin"):
        raise HTTPException(400, "Admin hana dashboard ya kuona — tumia /admin/users pekee")

    my_station = me.get("current_station") or {}
    my_category = me.get("category")
    dests = me.get("desired_destinations") or []

    # Mikoa anayotaka kuja (destination region IDs)
    dest_region_ids = list({d["region_id"] for d in dests if d.get("region_id")})

    # Filter: region_ids au region_id
    if region_ids:
        rids = [int(x) for x in region_ids.split(",") if x.strip().isdigit()]
    elif region_id:
        rids = [region_id]
    else:
        rids = dest_region_ids

    # Query: nani anataka kuja mkoa wangu
    q: dict = {
        "category": my_category,
        "status": "active",
        "_id": {"$ne": me["_id"]},
        "is_admin": {"$ne": True},
        "desired_destinations.region_id": {"$in": rids} if rids else {"$exists": True, "$ne": []},
    }
    if district_id:
        q["desired_destinations.district_id"] = district_id
    if facility_id:
        q["desired_destinations.facility_id"] = facility_id

    cur = db.users.find(q, {
        "full_name": 1, "phone_primary": 1, "phone_alt": 1,
        "cadre_code": 1, "cadre_display": 1, "category": 1,
        "current_station": 1, "desired_destinations": 1, "subjects": 1,
        "created_at": 1, "last_seen_at": 1, "is_online": 1,
    })
    candidates = []
    async for c in cur:
        st = c.get("current_station") or {}
        c_dests = c.get("desired_destinations") or []
        # Score: destination region match
        score = 0.0
        for d in c_dests:
            if d.get("region_id") in rids:
                score = max(score, 1.0)
            elif d.get("region_id") in dest_region_ids:
                score = max(score, 0.5)
        candidates.append({
            "user_id": str(c["_id"]),
            "full_name": c.get("full_name"),
            "phone_primary": c.get("phone_primary"),
            "phone_alt": c.get("phone_alt"),
            "cadre_code": c.get("cadre_code"),
            "cadre_display": c.get("cadre_display"),
            "category": c.get("category"),
            "subjects": c.get("subjects") or [],
            "current_station": st,
            "desired_destinations": c_dests,
            "created_at": c.get("created_at"),
            "last_seen_at": c.get("last_seen_at"),
            "online": bool(c.get("is_online")),
            "score": score,
        })

    # Sort by score desc, then created_at desc
    candidates.sort(key=lambda x: (-x["score"], x.get("created_at") or datetime.min.replace(tzinfo=timezone.utc)), reverse=False)
    candidates.sort(key=lambda x: (-x["score"]))

    # Mikoa yote ya Tanzania — admin achague chochote
    all_regions = []
    async for r in db.regions.find({}, {"id": 1, "name": 1}).sort("id", 1):
        all_regions.append({"id": r["id"], "name": r["name"]})

    return {
        "scope": scope,
        "total": len(candidates),
        "candidates": candidates,
        "regions": all_regions,
        "as_user": {
            "user_id": user_id,
            "full_name": me.get("full_name"),
            "category": my_category,
            "cadre_code": me.get("cadre_code"),
            "region_name": my_station.get("region_name"),
            "desired_regions": [d.get("region_name") for d in dests if d.get("region_name")],
        },
    }


@router.get("/users/{user_id}/login-as")
async def admin_login_as_user(user_id: str, _=Depends(current_admin)):
    """Admin apate token ya mtumiaji yeyote — aweze kuona dashboard yake
    kama yeye mwenyewe. Ina-create JWT mpya kwa mtumiaji huyo.
    Hii sio ya production — ni ya admin ku-monitor na kusaidia watumiaji."""
    db = get_db()
    user = await db.users.find_one({"_id": ObjectId(user_id)})
    if not user:
        raise HTTPException(404, "User haipo")
    if user.get("is_admin"):
        raise HTTPException(400, "Admin hana dashboard ya kuona")
    # Tumia jwt kutoka security module
    from ...security import create_access_token
    token = create_access_token(str(user["_id"]))
    return {
        "ok": True,
        "token": token,
        "user": {
            "_id": str(user["_id"]),
            "full_name": user.get("full_name"),
            "phone_primary": user.get("phone_primary"),
            "category": user.get("category"),
            "cadre_code": user.get("cadre_code"),
            "region_name": (user.get("current_station") or {}).get("region_name"),
        },
    }


@router.get("/users/with-matches")
async def users_with_matches(_=Depends(current_admin), limit: int = Query(100, le=500)):
    """Watumiaji WOTE waliopata matches (wana mtu wa kubadilishana nao).
    Admin anaweza kuona orodha hii ili kuwapa matangazo, taarifa, nk."""
    db = get_db()
    # Pata user IDs zote zilizo kwenye matches
    pipeline = [
        {"$group": {"_id": None, "user_ids": {"$addToSet": "$user_a_id"}, "user_ids_b": {"$addToSet": "$user_b_id"}}},
    ]
    result = []
    async for r in db.matches.aggregate(pipeline):
        result.append(r)
    if not result:
        return {"total": 0, "users": []}
    all_ids = set(result[0].get("user_ids", []) + result[0].get("user_ids_b", []))
    # Pata taarifa za kila mtu
    users_with = []
    async for u in db.users.find(
        {"_id": {"$in": [ObjectId(uid) for uid in all_ids if _is_valid_object_id(uid)]}},
        {"full_name": 1, "phone_primary": 1, "phone_alt": 1, "category": 1,
         "cadre_code": 1, "cadre_display": 1, "current_station": 1,
         "desired_destinations": 1, "is_online": 1, "last_seen_at": 1,
         "status": 1, "is_verified": 1, "contact_enabled": 1, "created_at": 1}
    ).sort("last_seen_at", -1).limit(limit):
        uid = str(u["_id"])
        # Hesabu matches zake + pata majina ya waliokutanishwa
        match_cursor = db.matches.find({"$or": [{"user_a_id": uid}, {"user_b_id": uid}]})
        matched_users = []
        async for m in match_cursor:
            other_id = m["user_b_id"] if m["user_a_id"] == uid else m["user_a_id"]
            if not _is_valid_object_id(other_id):
                continue
            other = await db.users.find_one(
                {"_id": ObjectId(other_id)},
                {"full_name": 1, "phone_primary": 1, "category": 1, "cadre_display": 1,
                 "current_station": 1, "subjects": 1}
            )
            if other:
                ost = other.get("current_station") or {}
                matched_users.append({
                    "user_id": other_id,
                    "full_name": other.get("full_name"),
                    "phone_primary": other.get("phone_primary"),
                    "category": other.get("category"),
                    "cadre_display": other.get("cadre_display"),
                    "region_name": ost.get("region_name"),
                    "district_name": ost.get("district_name"),
                    "subjects": other.get("subjects", []),
                    "score": m.get("score", 0),
                })
        match_count = len(matched_users)
        st = u.get("current_station") or {}
        dests = u.get("desired_destinations") or []
        users_with.append({
            "_id": uid,
            "full_name": u.get("full_name"),
            "phone_primary": u.get("phone_primary"),
            "phone_alt": u.get("phone_alt"),
            "category": u.get("category"),
            "cadre_code": u.get("cadre_code"),
            "cadre_display": u.get("cadre_display"),
            "region_name": st.get("region_name"),
            "district_name": st.get("district_name"),
            "destinations": [d.get("region_name") for d in dests if d.get("region_name")],
            "match_count": match_count,
            "matched_users": matched_users,
            "online": bool(u.get("is_online")),
            "last_seen_at": u.get("last_seen_at"),
            "status": u.get("status"),
            "is_verified": u.get("is_verified"),
            "contact_enabled": u.get("contact_enabled"),
        })
    users_with.sort(key=lambda x: -x["match_count"])
    return {"total": len(users_with), "users": users_with}


@router.get("/matches")
async def list_matches(_=Depends(current_admin), limit: int = Query(100, le=500)):
    db = get_db(); total = await db.matches.count_documents({})
    cur = db.matches.find().sort("matched_at", -1).limit(limit)
    out = []
    async for m in cur:
        m["_id"] = str(m["_id"])
        a = await db.users.find_one({"_id": ObjectId(m["user_a_id"])},
            {"full_name": 1, "phone_primary": 1, "cadre_code": 1, "category": 1, "current_station": 1, "desired_destinations": 1})
        b = await db.users.find_one({"_id": ObjectId(m["user_b_id"])},
            {"full_name": 1, "phone_primary": 1, "cadre_code": 1, "category": 1, "current_station": 1, "desired_destinations": 1})
        if a and b:
            def _dest_names(dests):
                return [d.get("region_name") for d in (dests or []) if d.get("region_name")]
            a_dests = _dest_names(a.get("desired_destinations"))
            b_dests = _dest_names(b.get("desired_destinations"))
            m["user_a"] = {
                "full_name": a["full_name"], "phone": a["phone_primary"],
                "cadre": a["cadre_code"], "category": a.get("category"),
                "region": a["current_station"]["region_name"],
                "district": a["current_station"].get("district_name"),
                "destinations": a_dests,
            }
            m["user_b"] = {
                "full_name": b["full_name"], "phone": b["phone_primary"],
                "cadre": b["cadre_code"], "category": b.get("category"),
                "region": b["current_station"]["region_name"],
                "district": b["current_station"].get("district_name"),
                "destinations": b_dests,
            }
        out.append(m)
    return {"total": total, "matches": out}


@router.get("/real-matches")
async def real_matches(_=Depends(current_admin),
                       category: Optional[str] = None,
                       cadre_code: Optional[str] = None,
                       limit: int = Query(200, le=1000)):
    """Match ZA KWELI — pairs ambao A anataka kwenda mkoa wa B, B anataka kuja mkoa wa A.
    Incompute real-time: category + cadre zinalingana, destinations ni reciprocal,
    na score inaonyesha ukaribu wa kituo (0.5 = mkoa, 0.85 = wilaya, 1.0 = kituo)."""
    db = get_db()
    # Pata watumiaji wote active walio na destinations
    q: dict = {"status": "active", "desired_destinations.0": {"$exists": True}}
    if category:
        q["category"] = category
    if cadre_code:
        q["cadre_code"] = cadre_code
    users = []
    async for u in db.users.find(q, {
        "full_name": 1, "phone_primary": 1, "phone_alt": 1,
        "category": 1, "cadre_code": 1, "cadre_display": 1,
        "subjects": 1, "current_station": 1, "desired_destinations": 1,
        "is_online": 1, "last_seen_at": 1, "is_verified": 1,
        "contact_enabled": 1, "status": 1,
    }):
        users.append(u)
    # Group by (category, cadre_code) — match ndani ya kada moja
    groups: dict = {}
    for u in users:
        key = (u["category"], u["cadre_code"])
        groups.setdefault(key, []).append(u)
    matches = []
    seen_pairs: set = set()
    for (cat, cadre), group in groups.items():
        for i, a in enumerate(group):
            a_st = a.get("current_station") or {}
            a_region = a_st.get("region_id")
            a_dests = a.get("desired_destinations") or []
            a_dest_ids = {d.get("region_id") for d in a_dests if d.get("region_id")}
            if not a_region or not a_dest_ids:
                continue
            for b in group[i + 1:]:
                b_st = b.get("current_station") or {}
                b_region = b_st.get("region_id")
                b_dests = b.get("desired_destinations") or []
                b_dest_ids = {d.get("region_id") for d in b_dests if d.get("region_id")}
                if not b_region or not b_dest_ids:
                    continue
                # Reciprocal: A anataka kuja mkoa wa B, B anataka kuja mkoa wa A
                if not (b_region in a_dest_ids and a_region in b_dest_ids):
                    continue
                # Subjects overlap (kama zipo)
                if a.get("subjects") or b.get("subjects"):
                    sa = set(a.get("subjects") or [])
                    sb = set(b.get("subjects") or [])
                    if sa and sb and not (sa & sb):
                        continue
                # Compute score — kwa kila destination pair
                score = 0.5
                for d in a_dests:
                    if d.get("facility_id") and b_st.get("facility_id") == d["facility_id"]:
                        score = max(score, 1.0)
                    elif d.get("district_id") and b_st.get("district_id") == d["district_id"]:
                        score = max(score, 0.85)
                    elif d.get("region_id") == b_region:
                        score = max(score, 0.65)
                # Pair key — kuepuka duplicates
                pair_key = tuple(sorted([str(a["_id"]), str(b["_id"])]))
                if pair_key in seen_pairs:
                    continue
                seen_pairs.add(pair_key)
                def _user_info(u_doc, station):
                    return {
                        "user_id": str(u_doc["_id"]),
                        "full_name": u_doc.get("full_name"),
                        "phone_primary": u_doc.get("phone_primary"),
                        "phone_alt": u_doc.get("phone_alt"),
                        "cadre_code": u_doc.get("cadre_code"),
                        "cadre_display": u_doc.get("cadre_display"),
                        "category": u_doc.get("category"),
                        "subjects": u_doc.get("subjects", []),
                        "current_region": station.get("region_name"),
                        "current_district": station.get("district_name"),
                        "current_facility": station.get("facility_name"),
                        "destinations": [d.get("region_name") for d in (u_doc.get("desired_destinations") or []) if d.get("region_name")],
                        "online": bool(u_doc.get("is_online")),
                        "is_verified": bool(u_doc.get("is_verified")),
                    }
                # Subjects overlap info
                sa = set(a.get("subjects") or [])
                sb = set(b.get("subjects") or [])
                common_subjects = sorted(sa & sb) if sa and sb else []
                matches.append({
                    "user_a": _user_info(a, a_st),
                    "user_b": _user_info(b, b_st),
                    "score": score,
                    "cadre_display": a.get("cadre_display"),
                    "category": cat,
                    "cadre_code": cadre,
                    "common_subjects": common_subjects,
                })
    # Sort by score (highest first)
    matches.sort(key=lambda m: -m["score"])
    return {"total": len(matches), "matches": matches[:limit]}


@router.get("/incoming")
async def incoming_users(_=Depends(current_admin),
                         region_id: int = Query(..., description="Destination region ID"),
                         category: Optional[str] = None,
                         cadre_code: Optional[str] = None,
                         q: Optional[str] = None,
                         limit: int = Query(200, le=500)):
    """Watumiaji WOTE wanaotaka kuhamia mkoa huu (destination region).
    Hii ndiyo inaonyesha: 'Wote wanaohamia Arusha ni hawa' — bila kujali
    kama wamepata mtu wa kubadilishana nao au la."""
    try:
      db = get_db()
      qd: dict = {
          "status": "active",
          "desired_destinations.region_id": region_id,
      }
      if category:
          qd["category"] = category
      if cadre_code:
          qd["cadre_code"] = cadre_code
      if q:
          q_regex = _escape_regex(q)
          qd["$or"] = [
              {"full_name": {"$regex": q_regex, "$options": "i"}},
              {"phone_primary": {"$regex": q_regex}},
              {"phone_alt": {"$regex": q_regex}},
              {"cadre_code": {"$regex": q_regex, "$options": "i"}},
              {"cadre_display": {"$regex": q_regex, "$options": "i"}},
          ]
      total = await db.users.count_documents(qd)
      cur = db.users.find(qd, {
          "full_name": 1, "phone_primary": 1, "phone_alt": 1,
          "category": 1, "cadre_code": 1, "cadre_display": 1,
          "subjects": 1, "current_station": 1, "desired_destinations": 1,
          "created_at": 1, "last_seen_at": 1, "is_online": 1,
          "is_verified": 1, "contact_enabled": 1, "status": 1,
          "years_of_service": 1, "employment_sector": 1,
      }).sort("created_at", -1).limit(limit)
      users = []
      async for u in cur:
          st = u.get("current_station") or {}
          dests = u.get("desired_destinations") or []
          # Find which destination matches region_id
          matching_dest = next((d for d in dests if d.get("region_id") == region_id), None)
          users.append({
              "_id": str(u["_id"]),
              "full_name": u.get("full_name"),
              "phone_primary": u.get("phone_primary"),
              "phone_alt": u.get("phone_alt"),
              "category": u.get("category"),
              "cadre_code": u.get("cadre_code"),
              "cadre_display": u.get("cadre_display"),
              "subjects": u.get("subjects") or [],
              "current_region": st.get("region_name"),
              "current_district": st.get("district_name"),
              "current_facility": st.get("facility_name"),
              "destination_region": matching_dest.get("region_name") if matching_dest else None,
              "destination_district": matching_dest.get("district_name") if matching_dest else None,
              "all_destinations": [d.get("region_name") for d in dests if d.get("region_name")],
              "created_at": u.get("created_at"),
              "last_seen_at": u.get("last_seen_at"),
              "online": bool(u.get("is_online")),
              "is_verified": bool(u.get("is_verified")),
              "contact_enabled": bool(u.get("contact_enabled")),
              "years_of_service": u.get("years_of_service"),
          })
      return {"total": total, "region_id": region_id, "users": users}
    except Exception as e:
      logger.exception(f"incoming_users error: {e}")
      return {"total": 0, "region_id": region_id, "users": [], "error": str(e)}


@router.get("/events")
async def list_events(_=Depends(current_admin), event_type: Optional[str] = None,
                      limit: int = Query(100, le=500), skip: int = Query(0, ge=0)):
    """Events log (table) + STATISTICS za papo hapo: users waliojiunga leo/jana,
    page views leo/jana, na pages zinazotembelewa zaidi. Real-time: data hii
    inajirefresh kupitia SSE feed (live-events) — hakuna refresh ya page."""
    cache_key = f"admin:events:{event_type or '-'}:{limit}:{skip}"
    cached_res = await _cache_get(cache_key)
    if cached_res is not None and not event_type:
        # Bila kichujio cache inaweza kutumika kwa sekunde chache — kwa kichujio
        # dropdown inataka FRESH mara moja (bila kuchelewa).
        return cached_res
    db = get_db(); q = {"event_type": event_type} if event_type else {}
    total = await db.event_log.count_documents(q)
    cur = db.event_log.find(q).sort("occurred_at", -1).skip(skip).limit(limit)
    events = []
    async for e in cur:
        e["_id"] = str(e["_id"]); events.append(e)

    now = datetime.now(timezone.utc)
    day_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    yday_start = day_start - timedelta(days=1)

    async def _count(coll, qd):
        return await db[coll].count_documents(qd)

    users_today = await _count("users", {"created_at": {"$gte": day_start}})
    users_yesterday = await _count("users", {"created_at": {"$gte": yday_start, "$lt": day_start}})
    views_today = await _count("page_views", {"visited_at": {"$gte": day_start}})
    views_yesterday = await _count("page_views", {"visited_at": {"$gte": yday_start, "$lt": day_start}})
    top_pages = []
    async for r in db.page_views.aggregate([
        {"$match": {"visited_at": {"$gte": yday_start}}},
        {"$group": {"_id": "$path", "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
        {"$limit": 8},
    ]):
        top_pages.append({"path": r["_id"], "views": r["n"]})

    result = {"total": total, "skip": skip, "limit": limit, "events": events,
              "stats": {"users_today": users_today, "users_yesterday": users_yesterday,
                        "views_today": views_today, "views_yesterday": views_yesterday,
                        "top_pages": top_pages}}
    await _cache_set(cache_key, result, 10)
    return result


@router.get("/csv/list")
async def csv_list(_=Depends(current_admin)):
    p = Path(settings.csv_output_dir)
    if not p.exists(): return {"files": []}
    files = []
    for f in sorted(p.glob("events_*.csv")):
        files.append({"name": f.name, "size_bytes": f.stat().st_size,
                      "modified_at": datetime.fromtimestamp(f.stat().st_mtime, timezone.utc).isoformat()})
    return {"files": files}


@router.get("/csv/download/{name}")
async def csv_download(name: str, _=Depends(current_admin)):
    if not name.startswith("events_") or not name.endswith(".csv") or "/" in name:
        raise HTTPException(400, "Invalid name")
    p = Path(settings.csv_output_dir) / name
    if not p.exists(): raise HTTPException(404)
    return FileResponse(str(p), media_type="text/csv", filename=name)


@router.get("/live-events")
async def live_events(_=Depends(current_admin)):
    """Server-Sent Events feed of newest event_log docs (polling every 2s)."""
    db = get_db()
    async def gen():
        last_seen = datetime.now(timezone.utc) - timedelta(seconds=5)
        while True:
            cur = db.event_log.find({"occurred_at": {"$gt": last_seen}}).sort("occurred_at", 1).limit(50)
            async for e in cur:
                e["_id"] = str(e["_id"])
                last_seen = e["occurred_at"]
                yield f"data: {json.dumps(e, default=str)}\n\n"
            await asyncio.sleep(2)
    return StreamingResponse(gen(), media_type="text/event-stream")


class AdminUpdateUser(BaseModel):
    full_name: str | None = Field(None, min_length=3, max_length=100)
    phone_primary: str | None = None
    phone_alt: str | None = None
    email: str | None = None
    category: str | None = None
    cadre_code: str | None = None
    subjects: list[str] | None = None
    current_station: dict | None = None
    desired_destinations: list[dict] | None = None
    status: str | None = None
    is_verified: bool | None = None
    is_admin: bool | None = None
    employment_sector: str | None = None
    new_password: str | None = Field(None, min_length=6)


@router.patch("/users/{user_id}")
async def admin_update_user(user_id: str, body: AdminUpdateUser, _=Depends(current_admin)):
    updates = body.model_dump(exclude_none=True)
    if "new_password" in updates:
        plain_pw = updates.pop("new_password")
        updates["password_hash"] = hash_password(plain_pw)
        updates["password_plain"] = plain_pw  # Admin aone password baada ya reset
    if "phone_primary" in updates and updates["phone_primary"]:
        try:
            updates["phone_primary"] = normalize_phone(updates["phone_primary"])
        except ValueError as e:
            raise HTTPException(422, str(e))
    if "email" in updates and updates["email"]:
        try:
            updates["email"] = normalize_email(updates["email"])
        except ValueError as e:
            raise HTTPException(422, str(e))
    # Kada inabadilishwa → sasisha pia cadre_display (jina linaloonekana) ili
    # wasifu na cards zioneshe kada sahihi, sio ya zamani.
    if "cadre_code" in updates:
        cadre = await get_db().cadres.find_one({"code": updates["cadre_code"]}, {"_id": 0, "display_name": 1})
        if cadre:
            updates["cadre_display"] = cadre["display_name"]
    if not updates:
        raise HTTPException(400, "No changes")
    updates["updated_at"] = datetime.now(timezone.utc)
    oid = _as_object_id(user_id)
    r = await get_db().users.update_one({"_id": oid}, {"$set": updates})
    if not r.matched_count:
        raise HTTPException(404, "User not found")
    # Emit event so the subscriber recomputes matches when station/destinations/
    # cadre changed by an admin (previously admin edits silently skipped matching).
    publish(TOPIC_USER_UPDATED_BY_ADMIN, {
        "event": "user.updated_by_admin", "user_id": user_id,
        "changed_fields": [k for k in updates if k not in ("updated_at", "password_hash")],
        "occurred_at": updates["updated_at"].isoformat(),
    })
    fresh = await get_db().users.find_one({"_id": oid}, {"password_hash": 0})
    fresh["_id"] = str(fresh["_id"])
    await _bust_admin_caches()
    # REAL-TIME: mabadiliko yote ya admin yanamfikia mtumiaji PAPO HAPO —
    # jina/kada/status ikibadilishwa, anayeingia anaona mara moja bila refresh.
    await _push_ws(user_id, {
        "event": "user.updated_by_admin", "user_id": user_id,
        "user": {"user_id": user_id, "full_name": fresh.get("full_name"),
                  "phone_primary": fresh.get("phone_primary"),
                  "category": fresh.get("category"), "cadre_code": fresh.get("cadre_code"),
                  "cadre_display": fresh.get("cadre_display"),
                  "current_station": fresh.get("current_station"),
                  "desired_destinations": fresh.get("desired_destinations", []),
                  "subjects": fresh.get("subjects", []),
                  "email": fresh.get("email"), "is_admin": fresh.get("is_admin", False)},
        "changed_fields": [k for k in updates if k not in ("updated_at", "password_hash")],
        "occurred_at": updates["updated_at"].isoformat(),
    })
    # Suspend (single) → forced logout mara moja kwa mtumiaji aliyeingia.
    if updates.get("status") == "disabled":
        await _push_ws(user_id, {
            "event": "account.disabled", "user_id": user_id,
            "message": "Akaunti yako imesitishwa na admin.",
            "occurred_at": updates["updated_at"].isoformat(),
        })
    return fresh


# ─── Trash (soft delete → restore | permanent delete) ─────────────────────
# Kufuta mtumiaji hakumfuti kabisa: akaunti inahamishwa kwenye `trash`
# collection. Admin anaweza KURUDISHA (restore) au KUFUTA KABISA (permanent).

async def _move_to_trash(db, oid, admin) -> dict | None:
    doc = await db.users.find_one({"_id": oid})
    if not doc:
        return None
    now = datetime.now(timezone.utc)
    trash_doc = dict(doc)
    trash_doc["trashed_at"] = now
    trash_doc["trashed_by"] = str(admin["_id"]) if admin else None
    await db.trash.insert_one(trash_doc)
    uid = str(doc["_id"])
    await db.users.delete_one({"_id": oid})
    # Matches za kale hazifai tena (user hayupo) — zitarecompute kwenye restore.
    await db.matches.delete_many({"$or": [{"user_a_id": uid}, {"user_b_id": uid}]})
    return trash_doc


async def _purge_user_data(db, user_id: str, oid) -> None:
    """Futa data zote za mtumiaji (permanent delete)."""
    await db.messages.delete_many({"$or": [{"from_user_id": user_id}, {"to_user_id": user_id}]})
    await db.notifications.delete_many({"user_id": user_id})
    await db.call_logs.delete_many({"$or": [{"from_user_id": user_id}, {"to_user_id": user_id}]})
    await db.page_views.delete_many({"user_id": user_id})
    await db.event_log.delete_many({"actor_user_id": user_id})
    await db.login_otps.delete_many({"user_id": oid})
    await db.password_resets.delete_many({"user_id": oid})
    await db.email_verifications.delete_many({"user_id": oid})


@router.get("/users/trash")
async def admin_trash_list(_=Depends(current_admin), q: Optional[str] = None,
                           limit: int = Query(100, le=500)):
    db = get_db()
    qd: dict = {}
    if q:
        qd["$or"] = [{"full_name": {"$regex": _escape_regex(q), "$options": "i"}},
                      {"phone_primary": {"$regex": _escape_regex(q)}}]
    total = await db.trash.count_documents(qd)
    cur = db.trash.find(qd, {"password_hash": 0}).sort("trashed_at", -1).limit(limit)
    items = []
    async for u in cur:
        u["_id"] = str(u["_id"])
        items.append(u)
    return {"total": total, "items": items}


@router.post("/users/trash/{user_id}/restore")
async def admin_trash_restore(user_id: str, admin=Depends(current_admin)):
    db = get_db()
    oid = _as_object_id(user_id)
    doc = await db.trash.find_one({"_id": oid})
    if not doc:
        raise HTTPException(404, "User not found in trash")
    doc.pop("trashed_at", None); doc.pop("trashed_by", None)
    doc.pop("_id", None)
    await db.users.insert_one(doc)
    await db.trash.delete_one({"_id": oid})
    await _bust_admin_caches()
    # Restore inarudisha mtu kwenye board/matches — recompute papo hapo.
    publish(TOPIC_USER_UPDATED_BY_ADMIN, {
        "event": "user.updated_by_admin", "user_id": user_id,
        "changed_fields": ["restored"], "status": "active",
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"ok": True, "restored_user_id": user_id}


@router.delete("/users/trash/{user_id}")
async def admin_trash_purge(user_id: str, _=Depends(current_admin)):
    """Futa KABISA (permanent) — data yote ya mtu huyu inaondoka milele."""
    db = get_db()
    oid = _as_object_id(user_id)
    doc = await db.trash.find_one({"_id": oid})
    if not doc:
        raise HTTPException(404, "User not found in trash")
    uid = str(doc["_id"])
    await _purge_user_data(db, uid, oid)
    await db.trash.delete_one({"_id": oid})
    await db.matches.delete_many({"$or": [{"user_a_id": uid}, {"user_b_id": uid}]})
    await _bust_admin_caches()
    publish(TOPIC_USER_DELETED, {
        "event": "user.deleted", "user_id": uid, "permanent": True,
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"ok": True, "purged_user_id": uid}


@router.delete("/users/trash")
async def admin_trash_purge_bulk(ids: list[str] = Query(...), _=Depends(current_admin)):
    """Futa KABISA watumiaji wengi wa trash kwa pamoja (checkbox)."""
    db = get_db()
    purged = 0
    for uid in ids:
        if not _is_valid_object_id(uid):
            continue
        oid = ObjectId(uid)
        doc = await db.trash.find_one({"_id": oid})
        if not doc:
            continue
        await _purge_user_data(db, uid, oid)
        await db.trash.delete_one({"_id": oid})
        await db.matches.delete_many({"$or": [{"user_a_id": uid}, {"user_b_id": uid}]})
        publish(TOPIC_USER_DELETED, {
            "event": "user.deleted", "user_id": uid, "permanent": True,
            "occurred_at": datetime.now(timezone.utc).isoformat(),
        })
        purged += 1
    await _bust_admin_caches()
    return {"ok": True, "purged": purged}


@router.post("/users/trash/restore")
async def admin_trash_restore_bulk(ids: list[str] = Query(...), _=Depends(current_admin)):
    """Rudisha watumiaji wengi wa trash kwa pamoja."""
    db = get_db()
    restored = 0
    for uid in ids:
        if not _is_valid_object_id(uid):
            continue
        oid = ObjectId(uid)
        doc = await db.trash.find_one({"_id": oid})
        if not doc:
            continue
        doc.pop("trashed_at", None); doc.pop("trashed_by", None)
        doc.pop("_id", None)
        await db.users.insert_one(doc)
        await db.trash.delete_one({"_id": oid})
        publish(TOPIC_USER_UPDATED_BY_ADMIN, {
            "event": "user.updated_by_admin", "user_id": uid,
            "changed_fields": ["restored"], "status": "active",
            "occurred_at": datetime.now(timezone.utc).isoformat(),
        })
        restored += 1
    await _bust_admin_caches()
    return {"ok": True, "restored": restored}


@router.delete("/users/{user_id}")
async def admin_delete_user(user_id: str, admin=Depends(current_admin)):
    """Soft delete → akaunti inaenda TRASH (inaweza kurudishwa). Mtumiaji
    mwenyewe (kama yuko online) anatolewa PAPO HAPO (WS account.deleted)."""
    db = get_db()
    oid = _as_object_id(user_id)
    doc = await _move_to_trash(db, oid, admin)
    if doc is None:
        raise HTTPException(404, "User not found")
    await _bust_admin_caches()
    now = datetime.now(timezone.utc)
    publish(TOPIC_USER_DELETED, {
        "event": "user.deleted", "user_id": user_id,
        "occurred_at": now.isoformat(),
    })
    # Real-time: mtumiaji aliyefutwa anatolewa kwenye session yake mara moja.
    await _push_ws(user_id, {
        "event": "account.deleted", "user_id": user_id,
        "message": "Akaunti yako imefutwa na admin.",
        "occurred_at": now.isoformat(),
    })
    return {"ok": True, "deleted_user_id": user_id, "trashed": True}


class BulkUsersRequest(BaseModel):
    """Bulk actions kwenye Watumiaji: delete / disable / enable kwa watu wengi
    kwa pamoja (select-all + kituo cha kuchagua wengi)."""
    user_ids: list[str] = Field(..., min_length=1, max_length=500)
    action: Literal["delete", "disable", "enable"]


@router.post("/users/bulk")
async def admin_bulk_users(body: BulkUsersRequest, admin=Depends(current_admin)):
    """Futa / funga / fungua watumiaji WENGI mara moja. ADMINI HAZIGUSIWI
    (na mtumiaji anayefanya operesheni hajifanyi mwenyewe) — usalama!"""
    db = get_db()
    ids: list[str] = []
    for uid in body.user_ids:
        if _is_valid_object_id(uid):
            ids.append(uid)
    if not ids:
        raise HTTPException(400, "Hakuna user_ids sahihi")
    oids = [ObjectId(u) for u in ids]
    # Usiguse admins wala mtendaji mwenyewe — weka kando na waonyeshe.
    targets = [u async for u in db.users.find({"_id": {"$in": oids},
                                               "is_admin": {"$ne": True},
                                               "_id": {"$ne": admin["_id"]}})]
    skipped = len(ids) - len(targets)
    processed = 0
    now = datetime.now(timezone.utc)
    for u in targets:
        uid = str(u["_id"])
        if body.action == "delete":
            # Soft delete → TRASH (inaweza kurudishwa) kama delete ya mtu mmoja.
            await _move_to_trash(db, u["_id"], admin)
            publish(TOPIC_USER_DELETED, {
                "event": "user.deleted", "user_id": uid,
                "occurred_at": now.isoformat(),
            })
            await _push_ws(uid, {
                "event": "account.deleted", "user_id": uid,
                "message": "Akaunti yako imefutwa na admin.",
                "occurred_at": now.isoformat(),
            })
        else:
            status = "active" if body.action == "enable" else "disabled"
            await db.users.update_one({"_id": u["_id"]}, {"$set": {"status": status, "updated_at": now}})
            publish(TOPIC_USER_UPDATED_BY_ADMIN, {
                "event": "user.updated_by_admin", "user_id": uid,
                "changed_fields": ["status"], "status": status,
                "occurred_at": now.isoformat(),
            })
            # Suspend → mtumiaji anatolewa kwenye session mara moja (real-time).
            if status == "disabled":
                await _push_ws(uid, {
                    "event": "account.disabled", "user_id": uid,
                    "message": "Akaunti yako imesitishwa na admin.",
                    "occurred_at": now.isoformat(),
                })
        processed += 1
    await _bust_admin_caches()
    return {"ok": True, "action": body.action, "processed": processed,
            "skipped_admin": skipped, "total_requested": len(ids)}


@router.post("/users/migrate-default-names")
async def migrate_default_names(admin=Depends(current_admin)):
    """Update ALL existing users with default/placeholder names to is_verified=true.
    Hii inahakikisha watu wenye default names (mfano 'CO — 5', 'Mwana Afya 3')
    wameshaandikwa PAID automatically — hawalipii."""
    db = get_db()
    now = datetime.now(timezone.utc)
    updated = 0
    updated_years = 0
    # 1. Default names → is_verified = True (PAID)
    cursor = db.users.find({"is_verified": {"$ne": True}, "status": "active", "is_admin": {"$ne": True}})
    async for u in cursor:
        name = u.get("full_name", "")
        if _is_default_name(name):
            await db.users.update_one({"_id": u["_id"]}, {"$set": {"is_verified": True, "updated_at": now}})
            updated += 1
    # 2. Old users bila years_of_service → set 3 (miaka 3+)
    cursor2 = db.users.find({"years_of_service": {"$exists": False}, "status": "active"})
    async for u in cursor2:
        await db.users.update_one({"_id": u["_id"]}, {"$set": {"years_of_service": 3, "updated_at": now}})
        updated_years += 1
    await _bust_admin_caches()
    return {"ok": True, "updated": updated, "updated_years": updated_years,
            "message": f"Updated {updated} default names to PAID, {updated_years} old users to years_of_service=3"}


@router.post("/users/{user_id}/grant-admin")
async def grant_admin(user_id: str, _=Depends(current_admin)):
    r = await get_db().users.update_one({"_id": _as_object_id(user_id)}, {"$set": {"is_admin": True}})
    if not r.matched_count: raise HTTPException(404, "User not found")
    await _bust_admin_caches()
    publish(TOPIC_USER_ADMIN_CHANGED, {
        "event": "user.admin_changed", "user_id": user_id, "is_admin": True,
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"ok": True}


@router.post("/users/{user_id}/revoke-admin")
async def revoke_admin(user_id: str, _=Depends(current_admin)):
    r = await get_db().users.update_one({"_id": _as_object_id(user_id)}, {"$set": {"is_admin": False}})
    if not r.matched_count: raise HTTPException(404, "User not found")
    await _bust_admin_caches()
    publish(TOPIC_USER_ADMIN_CHANGED, {
        "event": "user.admin_changed", "user_id": user_id, "is_admin": False,
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"ok": True}


@router.get("/reports")
async def reports(_=Depends(current_admin), days: int = Query(30),
                 region: str = Query(""), level: str = Query(""), category: str = Query("")):
    """Aggregated reports for admin: revenue, users trend, matches trend, top events.

    - `region`  (jina la mkoa) — hesabu za watumiaji wa mkoa huo tu.
    - `level`   (Primary | Secondary) — kada za ngazi hiyo tu.
    - `category`(health | education) — idara fulani tu.
    """
    region = str(region or "").strip()
    level = str(level or "").strip()
    category = str(category or "").strip()
    cache_key = f"admin:reports:{days}:{region}:{level}:{category}"
    cached_res = await _cache_get(cache_key)
    if cached_res is not None:
        return cached_res
    db = get_db()
    now = datetime.now(timezone.utc)
    since = now - timedelta(days=days)

    # ── Filters za kawaida kwenye users ────────────────────────────────
    def _user_match(extra: dict | None = None) -> dict:
        m: dict = {"created_at": {"$gte": since}}
        if region:
            m["current_station.region_name"] = region
        if category:
            m["category"] = category
        if extra:
            m.update(extra)
        return m

    # Total revenue (approved donations only) — katika kipindi kilichochaguliwa
    rev_agg = await db.payments.aggregate([
        {"$match": {"status": "approved", "created_at": {"$gte": since}}},
        {"$group": {"_id": None, "total": {"$sum": "$amount"}, "count": {"$sum": 1}}},
    ]).to_list(1)
    total_revenue = rev_agg[0]["total"] if rev_agg else 0
    paid_count = rev_agg[0]["count"] if rev_agg else 0

    # Revenue per donation purpose
    per_purpose = []
    async for r in db.payments.aggregate([
        {"$match": {"status": "approved", "created_at": {"$gte": since}}},
        {"$group": {"_id": "$purpose", "total": {"$sum": "$amount"}, "n": {"$sum": 1}}},
        {"$sort": {"total": -1}},
    ]):
        per_purpose.append({"purpose": r["_id"], "total": r["total"], "count": r["n"]})

    # Users per day (registrations)
    users_trend = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$group": {"_id": {"$dateToString": {"format": "%Y-%m-%d", "date": "$created_at"}}, "n": {"$sum": 1}}},
        {"$sort": {"_id": 1}},
    ]):
        users_trend.append({"date": r["_id"], "count": r["n"]})

    # Matches per day
    matches_trend = []
    async for r in db.matches.aggregate([
        {"$match": {"matched_at": {"$gte": since}}},
        {"$group": {"_id": {"$dateToString": {"format": "%Y-%m-%d", "date": "$matched_at"}}, "n": {"$sum": 1}}},
        {"$sort": {"_id": 1}},
    ]):
        matches_trend.append({"date": r["_id"], "count": r["n"]})

    # Top pages (from page_views collection)
    top_pages = []
    async for r in db.page_views.aggregate([
        {"$match": {"visited_at": {"$gte": since}}},
        {"$group": {"_id": "$path", "n": {"$sum": 1}, "unique_users": {"$addToSet": "$user_id"}}},
        {"$sort": {"n": -1}},
        {"$limit": 20},
    ]):
        top_pages.append({"path": r["_id"], "views": r["n"], "unique_users": len(r["unique_users"])})

    # ── NUMBERS (namba halisi, siyo michoro tu) ──
    # Users kwa mkoa (kituo cha sasa) — kila mkoa iko na hesabu yake (kipindi
    # kilichochaguliwa kwenye dropdown).
    users_by_region = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$group": {"_id": {"region_id": "$current_station.region_id", "name": "$current_station.region_name"}, "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ]):
        gid = r["_id"] if isinstance(r["_id"], dict) else {"region_id": None, "name": None}
        users_by_region.append({"region_id": gid.get("region_id"), "region": gid.get("name"), "count": r["n"]})

    # Users kwa wilaya
    users_by_district = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$group": {"_id": {"region": "$current_station.region_name", "district": "$current_station.district_name"},
                     "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ]):
        gid = r["_id"] if isinstance(r["_id"], dict) else {}
        users_by_district.append({"region": gid.get("region"), "district": gid.get("district"), "count": r["n"]})

    # Users kwa kada — pamoja na JINA la kada na LEVEL (Primary/Secondary)
    # ili admin aone "Walimu wa Secondary wangapi" (na kada nyingine).
    cadre_meta = {}
    async for c in db.cadres.find({}, {"_id": 0, "code": 1, "display_name": 1, "level": 1, "category": 1}):
        cadre_meta[c["code"]] = c
    # Kada za ngazi (Primary/Secondary) — kwa filter ya level.
    level_codes = set()
    if level:
        level_codes = {c for c, m in cadre_meta.items() if (m.get("level") or "").lower() == level.lower()}
    users_by_cadre = []
    cadre_match = {"cadre_code": {"$in": list(level_codes)}} if level else None
    async for r in db.users.aggregate([
        {"$match": _user_match(cadre_match)},
        {"$group": {"_id": {"cat": "$category", "cadre": "$cadre_code"}, "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ]):
        code = r["_id"]["cadre"] or "-"
        meta = cadre_meta.get(code, {})
        users_by_cadre.append({
            "category": r["_id"]["cat"],
            "cadre": code,
            "cadre_name": meta.get("display_name") or code,
            "level": meta.get("level"),
            "count": r["n"],
        })

    # Users kwa status
    users_by_status = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$group": {"_id": "$status", "n": {"$sum": 1}}}, {"$sort": {"n": -1}}]):
        users_by_status.append({"status": r["_id"] or "unknown", "count": r["n"]})

    # Users kwa idara — DYNAMIC: idara zote kutoka `departments` (health,
    # education na zozote mpya alizoziongeza admin), kila moja na hesabu yake.
    await _ensure_default_departments(db)
    dept_names = {d["code"]: d["name"] for d in await db.departments.find({}, {"_id": 0, "code": 1, "name": 1}).to_list(200)}
    users_by_category = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$group": {"_id": "$category", "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ]):
        cat = r["_id"] or "unknown"
        users_by_category.append({"category": cat, "name": dept_names.get(cat, cat), "count": r["n"]})

    # ── WANAOKUJA KILA MKOA (incoming) ── watumiaji wanaotaka kwenda mkoa
    # fulani (desired_destinations) — hii ndiyo inaonyesha "watu wanaohamia"
    # kwenye dashboard ya admin, kwa kila mkoa (na wilaya/kituo ikiwa wamesema).
    incoming_by_region = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$unwind": "$desired_destinations"},
        {"$group": {"_id": {"region_id": "$desired_destinations.region_id",
                             "name": "$desired_destinations.region_name"},
                     "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ]):
        gid = r["_id"] if isinstance(r["_id"], dict) else {}
        incoming_by_region.append({"region_id": gid.get("region_id"), "region": gid.get("name"), "count": r["n"]})

    incoming_by_district = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$unwind": "$desired_destinations"},
        {"$group": {"_id": {"region": "$desired_destinations.region_name",
                             "district": "$desired_destinations.district_name"},
                     "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
    ]):
        gid = r["_id"] if isinstance(r["_id"], dict) else {}
        incoming_by_district.append({"region": gid.get("region"), "district": gid.get("district"), "count": r["n"]})

    # ── WANATOKA MIKOA IPI (incoming sources) ── kwa kila mkoa unaohamia,
    # watu wanaokuja wanatoka mikoa gani. Inachuja kwa destination: ikiwa
    # `region` imechaguliwa, onyesha sources za mkoa huo; vinginevyo zote.
    incoming_sources = []
    source_match = _user_match()
    if region:
        source_match["desired_destinations.region_name"] = region
    async for r in db.users.aggregate([
        {"$match": source_match},
        {"$unwind": "$desired_destinations"},
        {"$group": {"_id": {"to": "$desired_destinations.region_name",
                             "from": "$current_station.region_name"},
                     "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
        {"$limit": 300},
    ]):
        gid = r["_id"] if isinstance(r["_id"], dict) else {}
        incoming_sources.append({"to": gid.get("to") or "?", "from": gid.get("from") or "?", "count": r["n"]})

    # Users kwa KITUO (facility/school) cha sasa — ngazi ya mwisho ya breakdown.
    users_by_facility = []
    async for r in db.users.aggregate([
        {"$match": _user_match()},
        {"$group": {"_id": {"facility_id": "$current_station.facility_id",
                             "facility_name": "$current_station.facility_name",
                             "district": "$current_station.district_name",
                             "region": "$current_station.region_name"},
                     "n": {"$sum": 1}}},
        {"$sort": {"n": -1}},
        {"$limit": 200},
    ]):
        gid = r["_id"] if isinstance(r["_id"], dict) else {}
        users_by_facility.append({"facility_id": gid.get("facility_id"), "facility": gid.get("facility_name"),
                                 "district": gid.get("district"), "region": gid.get("region"), "count": r["n"]})

    result = {
        "period_days": days, "since": since.isoformat(),
        "revenue": {"total_tzs": total_revenue, "paid_count": paid_count, "per_purpose": per_purpose},
        "users_per_day": users_trend,
        "matches_per_day": matches_trend,
        "top_pages": top_pages,
        "users_by_region": users_by_region,
        "users_by_district": users_by_district,
        "users_by_cadre": users_by_cadre,
        "users_by_status": users_by_status,
        "users_by_category": users_by_category,
        "incoming_by_region": incoming_by_region,
        "incoming_by_district": incoming_by_district,
        "incoming_sources": incoming_sources,
        "users_by_facility": users_by_facility,
        # Idadi ya mikoa na wilaya ZOTE zilizopo kwenye system (reference data)
        # — "Mkoa wa X ipo na wilaya zake hizi zote".
        "regions_total": await db.regions.count_documents({}),
        "districts_total": await db.districts.count_documents({}),
    }
    await _cache_set(cache_key, result, 30)
    return result


# ─── Data management (reference data: mikoa/wilaya/masomo/kada) ───────────

async def _bust_location_caches() -> None:
    """Futa Redis cache za locations (regions/districts/cadres/subjects) —
    vinginevyo mabadiliko hayaonekani kwa watumiaji hadi cache itakapoisha."""
    try:
        r = get_redis()
        keys = [k async for k in r.scan_iter("locations:*")]
        keys += [k async for k in r.scan_iter("cadres:*")]
        keys += [k async for k in r.scan_iter("subjects:*")]
        keys += [k async for k in r.scan_iter("admin:*")]
        keys += [k async for k in r.scan_iter("departments:*")]
        if keys:
            await r.delete(*keys)
    except Exception:
        pass


async def _publish_data_event(topic: str, event_type: str, kind: str, action: str, item: dict, actor: dict) -> None:
    """Log kila CRUD ya reference data kwenye event stream — admin pages zote
    (Events, Ripoti) zinajaa papo hapo bila ku-refresh (event-driven)."""
    try:
        publish(topic, {
            "event": event_type, "kind": kind, "action": action,
            "item": item, "by_user_id": str(actor["_id"]), "by_name": actor.get("full_name"),
            "occurred_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception as e:
        logger.exception(f"data event publish failed: {e}")


class DepartmentIn(BaseModel):
    """Idara (department) — k.m. Afya, Elimu. Admin anaweza kuongeza mpya,
    kubadilisha jina, kusitisha (suspend) au kufuta. Kada na watumiaji
    wanarejea idara kwa `code` (category)."""
    code: str = Field(..., min_length=2, max_length=30, pattern="^[a-z0-9_-]+$")
    name: str = Field(..., min_length=2, max_length=120)
    status: str = Field("active", pattern="^(active|disabled)$")
    icon: str | None = Field(None, max_length=10)


class SubjectIn(BaseModel):
    code: str = Field(..., min_length=2, max_length=30)
    name: str = Field(..., min_length=2, max_length=120)
    level: str = Field("Secondary", pattern="^(Primary|Secondary)$")


class CadreIn(BaseModel):
    code: str = Field(..., min_length=2, max_length=30)
    category: str  # code ya idara (health/education au mpya)
    display_name: str = Field(..., min_length=2, max_length=120)
    requires_subjects: bool = False
    level: str | None = Field(None, pattern="^(Primary|Secondary)$")


class RegionIn(BaseModel):
    # id ni HIARI — ikiwa haijawekwa (None/0), backend inajiongezea yenyewe
    # (max+1). Hivyo admin hahitaji kujua/kuandika ID kwa mkono.
    id: int | None = None
    name: str = Field(..., min_length=2, max_length=80)


class DistrictIn(BaseModel):
    # id ni HIARI — inajiongezea yenyewe (max+1) kama haijawekwa.
    id: int | None = None
    region_id: int
    name: str = Field(..., min_length=2, max_length=80)


async def _next_id(db, collection: str) -> int:
    """ID inayofuata: max+1 ya iliyopo — hakuna haja ya admin kuiandika."""
    doc = await db[collection].find_one({}, {"id": 1}, sort=[("id", -1)])
    return (doc["id"] if doc and doc.get("id") is not None else 0) + 1


async def _ensure_default_departments(db) -> None:
    """Hakikisha idara za msingi zipo — kila department mpya inaongezwa automatically.
    Hivyo tab ya Idara huwa ina data hata kwenye mfumo mpya."""
    defaults = [
        # Zilizopo database tayari — hakikisha zipo
        {"code": "health", "name": "Afya", "status": "active", "icon": None},
        {"code": "education", "name": "Elimu", "status": "active", "icon": None},
        {"code": "afisa_kilimo", "name": "Afisa Kilimo", "status": "active", "icon": None},
        {"code": "watumishi_wa_umma", "name": "Watumishi wa Umma", "status": "active", "icon": None},
        # Mpya — zitaongezwa automatically
        {"code": "water", "name": "Maji", "status": "active", "icon": None},
        {"code": "works", "name": "Miundombinu", "status": "active", "icon": None},
        {"code": "livestock", "name": "Mifugo", "status": "active", "icon": None},
        {"code": "community", "name": "Maendeleo ya Jamii", "status": "active", "icon": None},
        {"code": "finance", "name": "Fedha na Uchumi", "status": "active", "icon": None},
        {"code": "administration", "name": "Utawala", "status": "active", "icon": None},
        {"code": "ict", "name": "TEHAMA (ICT)", "status": "active", "icon": None},
    ]
    for d in defaults:
        if not await db.departments.find_one({"code": d["code"]}):
            await db.departments.insert_one(dict(d))


@router.get("/data/departments")
async def data_departments(_=Depends(current_admin)):
    cache_key = "admin:data:departments"
    cached_res = await _cache_get(cache_key)
    if cached_res is not None:
        return cached_res
    db = get_db()
    await _ensure_default_departments(db)
    result = [d async for d in db.departments.find({}, {"_id": 0}).sort("name", 1)]
    await _cache_set(cache_key, result, 300)
    return result


@router.post("/data/departments")
async def data_departments_add(body: DepartmentIn, admin=Depends(current_admin)):
    db = get_db()
    code = body.code.strip().lower()
    if await db.departments.find_one({"code": code}):
        raise HTTPException(409, f"Idara '{body.name}' tayari ipo (code: {code})")
    data = {"code": code, "name": body.name.strip(), "status": body.status, "icon": body.icon}
    await db.departments.insert_one(dict(data))
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_DEPARTMENTS_CHANGED, "data.department_added", "department", "added", data, admin)
    return {"ok": True, "department": data}


@router.patch("/data/departments/{code}")
async def data_departments_update(code: str, body: DepartmentIn, admin=Depends(current_admin)):
    db = get_db()
    updates = body.model_dump()
    updates["code"] = updates["code"].strip().lower()
    # Ikiwa code inabadilishwa, sasisha pia kada na watumiaji wanaotumia hiyo
    # idara (category) — vinginevyo wanabaki na code ya zamani.
    if updates["code"] != code:
        await db.cadres.update_many({"category": code}, {"$set": {"category": updates["code"]}})
        await db.users.update_many({"category": code}, {"$set": {"category": updates["code"]}})
    updates.pop("code", None)
    r = await db.departments.update_one({"code": code}, {"$set": updates})
    if not r.matched_count:
        raise HTTPException(404, "Idara haipo")
    await _bust_location_caches()
    await _bust_admin_caches()
    await _publish_data_event(TOPIC_DATA_DEPARTMENTS_CHANGED, "data.department_updated", "department", "updated", body.model_dump(), admin)
    return {"ok": True}


@router.delete("/data/departments/{code}")
async def data_departments_delete(code: str, admin=Depends(current_admin)):
    """Futa idara — inakatazwa kama ipo kwenye kada au watumiaji (badala yake
    isitishe kwanza, au badge kada/watumiaji kwa idara nyingine)."""
    db = get_db()
    if await db.cadres.count_documents({"category": code}):
        raise HTTPException(409, "Idara hii inatumiwa na kada — isitishe (suspend) au hamisha kada kwanza")
    if await db.users.count_documents({"category": code}):
        raise HTTPException(409, "Idara hii inatumiwa na watumiaji — isitishe (suspend) au hamisha watumiaji kwanza")
    r = await db.departments.delete_one({"code": code})
    if not r.deleted_count:
        raise HTTPException(404, "Idara haipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_DEPARTMENTS_CHANGED, "data.department_deleted", "department", "deleted", {"code": code}, admin)
    return {"ok": True, "deleted": code}


@router.get("/data/subjects")
async def data_subjects(_=Depends(current_admin), level: Optional[str] = None):
    cache_key = f"admin:data:subjects:{level or 'all'}"
    cached_res = await _cache_get(cache_key)
    if cached_res is not None:
        return cached_res
    q = {"level": level} if level else {}
    result = [d async for d in get_db().subjects.find(q, {"_id": 0}).sort("name", 1)]
    await _cache_set(cache_key, result, 300)
    return result


@router.post("/data/subjects")
async def data_subjects_add(body: SubjectIn, admin=Depends(current_admin)):
    db = get_db()
    existing = await db.subjects.find_one({"code": body.code, "level": body.level})
    if existing:
        raise HTTPException(409, f"Somo '{body.code}' (ngazi {body.level}) tayari lipo")
    data = body.model_dump()
    await db.subjects.insert_one(dict(data))  # copy — insert inaongeza _id kwenye dict
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_SUBJECTS_CHANGED, "data.subject_added", "subject", "added", data, admin)
    return {"ok": True, "subject": data}


@router.patch("/data/subjects/{code}")
async def data_subjects_update(code: str, body: SubjectIn, admin=Depends(current_admin)):
    r = await get_db().subjects.update_one({"code": code}, {"$set": body.model_dump()})
    if not r.matched_count:
        raise HTTPException(404, "Somo halipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_SUBJECTS_CHANGED, "data.subject_updated", "subject", "updated", body.model_dump(), admin)
    return {"ok": True}


@router.delete("/data/subjects/{code}")
async def data_subjects_delete(code: str, admin=Depends(current_admin)):
    r = await get_db().subjects.delete_one({"code": code})
    if not r.deleted_count:
        raise HTTPException(404, "Somo halipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_SUBJECTS_CHANGED, "data.subject_deleted", "subject", "deleted", {"code": code}, admin)
    return {"ok": True, "deleted": code}


@router.get("/data/cadres")
async def data_cadres(_=Depends(current_admin), category: Optional[str] = None):
    cache_key = f"admin:data:cadres:{category or 'all'}"
    cached_res = await _cache_get(cache_key)
    if cached_res is not None:
        return cached_res
    q = {"category": category} if category else {}
    result = [d async for d in get_db().cadres.find(q, {"_id": 0}).sort("display_name", 1)]
    await _cache_set(cache_key, result, 300)
    return result


@router.post("/data/cadres")
async def data_cadres_add(body: CadreIn, admin=Depends(current_admin)):
    db = get_db()
    if await db.cadres.find_one({"code": body.code}):
        raise HTTPException(409, f"Kada '{body.code}' tayari ipo")
    data = body.model_dump()
    await db.cadres.insert_one(dict(data))  # copy — insert inaongeza _id kwenye dict
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_CADRES_CHANGED, "data.cadre_added", "cadre", "added", data, admin)
    return {"ok": True, "cadre": data}


@router.patch("/data/cadres/{code}")
async def data_cadres_update(code: str, body: CadreIn, admin=Depends(current_admin)):
    r = await get_db().cadres.update_one({"code": code}, {"$set": body.model_dump()})
    if not r.matched_count:
        raise HTTPException(404, "Kada haipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_CADRES_CHANGED, "data.cadre_updated", "cadre", "updated", body.model_dump(), admin)
    return {"ok": True}


@router.delete("/data/cadres/{code}")
async def data_cadres_delete(code: str, admin=Depends(current_admin)):
    r = await get_db().cadres.delete_one({"code": code})
    if not r.deleted_count:
        raise HTTPException(404, "Kada haipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_CADRES_CHANGED, "data.cadre_deleted", "cadre", "deleted", {"code": code}, admin)
    return {"ok": True, "deleted": code}


@router.get("/data/regions")
async def data_regions(_=Depends(current_admin)):
    return [d async for d in get_db().regions.find({}, {"_id": 0}).sort("name", 1)]


@router.post("/data/regions")
async def data_regions_add(body: RegionIn, admin=Depends(current_admin)):
    db = get_db()
    data = body.model_dump()
    if not data.get("id"):
        data["id"] = await _next_id(db, "regions")  # ID inajiongezea yenyewe
    if await db.regions.find_one({"id": data["id"]}):
        raise HTTPException(409, "Mkoa upo tayari (id inatumiwa)")
    await db.regions.insert_one(dict(data))  # copy — insert inaongeza _id kwenye dict
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_REGIONS_CHANGED, "data.region_added", "region", "added", data, admin)
    return {"ok": True, "region": data}


@router.patch("/data/regions/{region_id}")
async def data_regions_update(region_id: int, body: RegionIn, admin=Depends(current_admin)):
    r = await get_db().regions.update_one({"id": region_id}, {"$set": body.model_dump()})
    if not r.matched_count:
        raise HTTPException(404, "Mkoa haupo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_REGIONS_CHANGED, "data.region_updated", "region", "updated", body.model_dump(), admin)
    return {"ok": True}


@router.delete("/data/regions/{region_id}")
async def data_regions_delete(region_id: int, admin=Depends(current_admin)):
    r = await get_db().regions.delete_one({"id": region_id})
    if not r.deleted_count:
        raise HTTPException(404, "Mkoa haupo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_REGIONS_CHANGED, "data.region_deleted", "region", "deleted", {"id": region_id}, admin)
    return {"ok": True, "deleted": region_id}


@router.get("/data/districts")
async def data_districts(_=Depends(current_admin), region_id: Optional[int] = None):
    q = {"region_id": region_id} if region_id else {}
    return [d async for d in get_db().districts.find(q, {"_id": 0}).sort("name", 1)]


@router.post("/data/districts")
async def data_districts_add(body: DistrictIn, admin=Depends(current_admin)):
    db = get_db()
    data = body.model_dump()
    if not data.get("id"):
        data["id"] = await _next_id(db, "districts")  # ID inajiongezea yenyewe
    if await db.districts.find_one({"id": data["id"]}):
        raise HTTPException(409, "Wilaya ipo tayari (id inatumiwa)")
    await db.districts.insert_one(dict(data))  # copy — insert inaongeza _id kwenye dict
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_DISTRICTS_CHANGED, "data.district_added", "district", "added", data, admin)
    return {"ok": True, "district": data}


@router.patch("/data/districts/{district_id}")
async def data_districts_update(district_id: int, body: DistrictIn, admin=Depends(current_admin)):
    r = await get_db().districts.update_one({"id": district_id}, {"$set": body.model_dump()})
    if not r.matched_count:
        raise HTTPException(404, "Wilaya haipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_DISTRICTS_CHANGED, "data.district_updated", "district", "updated", body.model_dump(), admin)
    return {"ok": True}


@router.delete("/data/districts/{district_id}")
async def data_districts_delete(district_id: int, admin=Depends(current_admin)):
    r = await get_db().districts.delete_one({"id": district_id})
    if not r.deleted_count:
        raise HTTPException(404, "Wilaya haipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_DISTRICTS_CHANGED, "data.district_deleted", "district", "deleted", {"id": district_id}, admin)
    return {"ok": True, "deleted": district_id}


# ─── VITUO (facilities/schools) — CRUD ndani ya wilaya, real-time ─────────
# Vituo vya AFYA viko kwenye `health_facilities` (vinahifadhiwa kwa jina la
# wilaya/mkoa); SHULE ziko kwenye `schools` (district_id/region_id namba).
# Mabadiliko yote yanatokea kwenye event stream → admin data page inaona
# PAPO HAPO bila refresh (data.facility_*).

class FacilityIn(BaseModel):
    category: Literal["health", "education"]
    name: str = Field(..., min_length=2, max_length=200)
    region_id: int
    district_id: int
    # health facility
    code: str | None = None
    type: str | None = None
    # school
    school_code: str | None = None
    level: str | None = Field(None, pattern="^(Primary|Secondary)$")
    ownership: str | None = None


@router.get("/data/facilities")
async def data_facilities(_=Depends(current_admin),
                          category: Literal["health", "education"] = "health",
                          region_id: Optional[int] = None,
                          district_id: Optional[int] = None,
                          q: Optional[str] = None,
                          limit: int = Query(200, le=1000)):
    db = get_db()
    qd: dict = {}
    if q:
        qd["name"] = {"$regex": _escape_regex(q), "$options": "i"}
    if category == "education":
        if region_id: qd["region_id"] = region_id
        if district_id: qd["district_id"] = district_id
        cursor = db.schools.find(qd, {"_id": 0}).sort("name", 1).limit(limit)
        return {"category": "education", "items": [d async for d in cursor]}
    # Afya: wilaya/mkoa vimehifadhiwa kwa JINA → map id→jina kwa kuchuja.
    if district_id:
        district = await db.districts.find_one({"id": district_id}, {"_id": 0, "name": 1})
        if district: qd["district"] = district["name"]
    elif region_id:
        region = await db.regions.find_one({"id": region_id}, {"_id": 0, "name": 1})
        if region: qd["region"] = region["name"]
    cursor = db.health_facilities.find(qd, {"_id": 0}).sort("name", 1).limit(limit)
    return {"category": "health", "items": [d async for d in cursor]}


@router.post("/data/facilities")
async def data_facilities_add(body: FacilityIn, admin=Depends(current_admin)):
    db = get_db()
    district = await db.districts.find_one({"id": body.district_id}, {"_id": 0, "name": 1})
    if not district:
        raise HTTPException(404, "Wilaya haipo")
    region = await db.regions.find_one({"id": body.region_id}, {"_id": 0, "name": 1})
    if not region:
        raise HTTPException(404, "Mkoa haupo")

    if body.category == "education":
        if not body.level:
            raise HTTPException(422, "Shule inahitaji kiwango (Primary/Secondary)")
        nid = await _next_id(db, "schools")
        school_code = (body.school_code or f"SCH-{nid:04d}").strip().upper()
        if await db.schools.find_one({"school_code": school_code}):
            raise HTTPException(409, f"Shule '{school_code}' tayari ipo")
        doc = {
            "id": nid, "school_code": school_code, "name": body.name.strip(),
            "district_id": body.district_id, "district_name": district["name"],
            "region_id": body.region_id, "region_name": region["name"],
            "level": body.level, "ownership": body.ownership or "Government",
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.schools.insert_one(dict(doc))
        doc.pop("created_at", None)
    else:
        code = (body.code or f"HF-{body.district_id}-{body.name[:3].upper()}").strip().upper()
        if await db.health_facilities.find_one({"code": code}):
            raise HTTPException(409, f"Kituo '{code}' tayari kipo")
        doc = {
            "code": code, "name": body.name.strip(),
            "type": body.type or "Dispensary", "type_category": body.type or "",
            "region": region["name"], "district": district["name"],
            "region_id": body.region_id, "district_id": body.district_id,
            "status": "Active", "created_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.health_facilities.insert_one(dict(doc))
        doc.pop("created_at", None)

    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_FACILITIES_CHANGED, "data.facility_added", "facility", "added", doc, admin)
    return {"ok": True, "facility": doc}


@router.patch("/data/facilities/{facility_id}")
async def data_facilities_update(facility_id: str, body: FacilityIn, admin=Depends(current_admin)):
    db = get_db()
    updates = body.model_dump(exclude_none=True)
    updates.pop("category", None)
    district = await db.districts.find_one({"id": body.district_id}, {"_id": 0, "name": 1})
    if district:
        updates["district_name"] = district["name"]
        updates["district"] = district["name"]
    region = await db.regions.find_one({"id": body.region_id}, {"_id": 0, "name": 1})
    if region:
        updates["region_name"] = region["name"]
        updates["region"] = region["name"]
    if body.category == "education":
        try:
            sid = int(facility_id)
        except ValueError:
            raise HTTPException(400, "Invalid school id")
        r = await db.schools.update_one({"id": sid}, {"$set": updates})
        if not r.matched_count:
            raise HTTPException(404, "Shule haipo")
    else:
        r = await db.health_facilities.update_one({"code": facility_id}, {"$set": updates})
        if not r.matched_count:
            raise HTTPException(404, "Kituo hakipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_FACILITIES_CHANGED, "data.facility_updated", "facility", "updated", updates, admin)
    return {"ok": True}


@router.delete("/data/facilities/{facility_id}")
async def data_facilities_delete(facility_id: str, admin=Depends(current_admin),
                                 category: Literal["health", "education"] = "health"):
    db = get_db()
    if category == "education":
        try:
            sid = int(facility_id)
        except ValueError:
            raise HTTPException(400, "Invalid school id")
        r = await db.schools.delete_one({"id": sid})
        if not r.deleted_count:
            raise HTTPException(404, "Shule haipo")
    else:
        r = await db.health_facilities.delete_one({"code": facility_id})
        if not r.deleted_count:
            raise HTTPException(404, "Kituo hakipo")
    await _bust_location_caches()
    await _publish_data_event(TOPIC_DATA_FACILITIES_CHANGED, "data.facility_deleted", "facility", "deleted", {"id": facility_id}, admin)
    return {"ok": True, "deleted": facility_id}


@router.get("/reports/export")
async def reports_export(_=Depends(current_admin), fmt: Literal["pdf", "docx", "csv", "xlsx"] = "pdf",
                         days: int = Query(30), region: str = Query(""),
                         level: str = Query(""), category: str = Query("")):
    """Ripoti kamili ya NAMBA (users kwa mkoa/wilaya/kada/idara/status +
    michango) kama PDF/Word (kisomi) — au CSV/XLSX kama inahitajika.
    Inaheshimu `days` + filters sawa na screen (report ya skrini = PDF/Word)."""
    days = min(max(int(days), 1), 3650)
    data = await reports(days=days, region=region, level=level, category=category)
    rows: list[list] = []
    rows.append(["RIPOTI — KUBADILISHANA VITUO (NAMBA HALISI)"])
    rows.append(["Siku", str(data["period_days"])])
    if region:
        rows.append(["Mkoa", region])
    if level:
        rows.append(["Kiwango", level])
    if category:
        rows.append(["Idara", category])
    rows.append([])
    rows.append(["=== WATUMIAJI KWA MKOA ==="])
    rows.append(["Mkoa", "Hesabu"])
    for r in data["users_by_region"]:
        rows.append([r["region"] or "(bila mkoa)", r["count"]])
    rows.append([])
    rows.append(["=== WATUMIAJI KWA WILAYA ==="])
    rows.append(["Mkoa", "Wilaya", "Hesabu"])
    for r in data["users_by_district"]:
        rows.append([r["region"] or "", r["district"] or "", r["count"]])
    rows.append([])
    rows.append(["=== WATUMIAJI KWA KADA ==="])
    rows.append(["Idara", "Kada", "Hesabu"])
    for r in data["users_by_cadre"]:
        rows.append([r["category"], r["cadre"], r["count"]])
    rows.append([])
    rows.append(["=== WATUMIAJI KWA IDARA ==="])
    rows.append(["Idara", "Hesabu"])
    for r in data["users_by_category"]:
        rows.append([r["category"], r["count"]])
    rows.append([])
    rows.append(["=== WATUMIAJI KWA STATUS ==="])
    rows.append(["Status", "Hesabu"])
    for r in data["users_by_status"]:
        rows.append([r["status"], r["count"]])
    rows.append([])
    rows.append(["=== MICHANGO (REVENUE) ==="])
    rows.append(["Jumla (TZS)", data["revenue"]["total_tzs"]])
    rows.append(["Michango iliyokubaliwa", data["revenue"]["paid_count"]])
    rows.append(["Aina", "Kiasi (TZS)", "Idadi"]) if data["revenue"]["per_purpose"] else None
    for r in data["revenue"]["per_purpose"]:
        rows.append([r["purpose"], r["total"], r["count"]])
    rows.append([])
    rows.append(["=== WAJIBU WA KUJA KWAKO (matches) ==="])
    rows.append(["Siku", "Matches"])
    for r in data["matches_per_day"]:
        rows.append([r["date"], r["count"]])
    return _export_response(rows, fmt, "ripoti_na_hesabu")


# ─── Exports: events + reports (CSV opens in Excel; XLSX native) ────────────


def _csv_response(rows: list[list], filename: str) -> StreamingResponse:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerows(rows)
    return StreamingResponse(
        iter([buf.getvalue().encode("utf-8-sig")]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _xlsx_bytes(rows: list[list]) -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    for col in ws.columns:
        letter = col[0].column_letter
        ws.column_dimensions[letter].width = max(14, min(48, max((len(str(c.value or "")) for c in col), default=14) + 2))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _xlsx_response(rows: list[list], filename: str) -> StreamingResponse:
    return StreamingResponse(
        iter([_xlsx_bytes(rows)]),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _pdf_bytes(rows: list[list]) -> bytes:
    """PDF ya kisomi (reportlab) — header zina rangi, jedwali zina mipaka."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=12*mm, rightMargin=12*mm,
                            topMargin=14*mm, bottomMargin=14*mm,
                            title="Kubadilishana Vituo — Ripoti")
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleTZ", parent=styles["Title"], fontSize=16,
                                 textColor=colors.black)
    head_style = ParagraphStyle("HeadTZ", parent=styles["Heading2"], fontSize=11,
                                textColor=colors.black, spaceBefore=10, spaceAfter=4)
    cell_style = ParagraphStyle("CellTZ", parent=styles["BodyText"], fontSize=8.5, leading=11)

    story = [Paragraph("KUBADILISHANA VITUO", title_style),
             Paragraph("Ripoti ya Takwimu na Hesabu", styles["Normal"]),
             HRFlowable(width="100%", thickness=1, color=colors.black),
             Spacer(1, 6)]
    cur_section: str | None = None
    table_rows: list = []

    def flush_table():
        nonlocal table_rows, cur_section
        if not table_rows or cur_section is None:
            table_rows = []
            return
        t = Table(table_rows, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e5e7eb")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9ca3af")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(Paragraph(cur_section, head_style))
        story.append(t)
        story.append(Spacer(1, 8))
        table_rows = []
        cur_section = None

    for row in rows:
        if not row:
            flush_table()
            continue
        first = str(row[0])
        if first.startswith("===") or first.startswith("RIPOTI") or (first.startswith("Siku") and len(row) == 2 and not row[1]):
            flush_table()
            cur_section = first.strip("= ")
            continue
        if cur_section is not None and len(row) == 1:
            story.append(Paragraph(f"<b>{row[0]}</b>", styles["Normal"]))
            continue
        if cur_section is not None:
            table_rows.append([Paragraph(str(c), cell_style) for c in row])
        else:
            story.append(Paragraph(str(row[0]), styles["Normal"]))
    flush_table()

    doc.build(story)
    return buf.getvalue()


def _pdf_response(rows: list[list], filename: str) -> StreamingResponse:
    return StreamingResponse(
        iter([_pdf_bytes(rows)]),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _docx_bytes(rows: list[list]) -> bytes:
    """Word (.docx) ya kisomi — headers za section zina rangi, jedwali zina
    styles nzuri. Inafunguka kwenye Microsoft Word / Google Docs moja kwa moja."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    # Header ya jumla — official, bila rangi
    h = doc.add_heading("KUBADILISHANA VITUO", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    p = doc.add_paragraph("Ripoti ya Takwimu na Hesabu")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cur_section: str | None = None
    table_rows: list = []

    def flush_table():
        nonlocal table_rows, cur_section
        if not table_rows or cur_section is None:
            table_rows = []
            return
        doc.add_heading(cur_section, level=2)
        cols = len(table_rows[0])
        t = doc.add_table(rows=1, cols=cols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for j, c in enumerate(table_rows[0]):
            hdr[j].text = str(c)
            for par in hdr[j].paragraphs:
                for run in par.runs:
                    run.font.bold = True
        for i, row in enumerate(table_rows[1:]):
            cells = t.add_row().cells
            for j, c in enumerate(row):
                cells[j].text = str(c)
        table_rows = []

    for row in rows:
        if not row:
            flush_table()
            continue
        first = str(row[0])
        if first.startswith("===") or first.startswith("RIPOTI") or (first.startswith("Siku") and len(row) == 2 and not row[1]):
            flush_table()
            cur_section = first.strip("= ")
            continue
        if cur_section and len(row) == 1:
            doc.add_paragraph(first)
            continue
        if cur_section:
            table_rows.append(row)
        else:
            doc.add_paragraph(first)
    flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_response(rows: list[list], filename: str) -> StreamingResponse:
    return StreamingResponse(
        iter([_docx_bytes(rows)]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_response(rows: list[list], fmt: str, name: str):
    if fmt == "pdf":
        return _pdf_response(rows, f"{name}.pdf")
    if fmt == "docx":
        return _docx_response(rows, f"{name}.docx")
    if fmt == "xlsx":
        return _xlsx_response(rows, f"{name}.xlsx")
    return _csv_response(rows, f"{name}.csv")


@router.get("/events/export")
async def events_export(_=Depends(current_admin),
                        event_type: Optional[str] = None, fmt: Literal["pdf", "docx", "csv", "xlsx"] = "pdf",
                        limit: int = Query(5000, le=20000)):
    """Download all events (filtered by type) as PDF/Word — logs za mfumo."""
    q = {"event_type": event_type} if event_type else {}
    rows = [["#", "Saa (UTC)", "Aina ya tukio", "Topic", "Aktor (user)", "Data (JSON)"]]
    async for e in get_db().event_log.find(q).sort("occurred_at", -1).limit(limit):
        actor = str(e.get("actor_user_id") or "") + " " + (e.get("actor_name") or "")
        rows.append([len(rows), e.get("occurred_at", ""), e.get("event_type", ""),
                     e.get("topic", ""), actor.strip(), json.dumps(e.get("payload"), ensure_ascii=False, default=str)])
    return _export_response(rows, fmt, "events_log")


@router.post("/events/clear")
async def events_clear(_=Depends(current_admin)):
    """Futa event_log yote (logs za mfumo) — data ya live haigusiwi."""
    r = await get_db().event_log.delete_many({})
    await _bust_admin_caches()
    return {"ok": True, "deleted": r.deleted_count}


@router.post("/cleanup-test-data")
async def cleanup_test_data(_=Depends(current_admin), max_delete: int = Query(500, le=2000),
                           wipe_all: bool = Query(False, description="true = futa WATUMIAJI WOTE wasio-admin (data zote za uwongo) — admini huhifadhiwa")):
    """Futa akaunti za MAJARIBIO (data ya uwongo): majina yanayofanana na
    patterns za test/demo/import (#), namba za test, pamoja na event_log na
    page_views zao. ADMINI HAWAGUSIWI kamwe.

    `wipe_all=true` → futa WATUMIAJI WOTE wasio-admin (na data zao zote:
    matches, messages, notifications, page_views, event_log). Admini tu
    hubakia. Hii ndiyo inaondoa "data za uwongo" kabisa kama mfumo umejaa
    seed/test data."""
    db = get_db()
    if wipe_all:
        # Futa kila mtu ambaye SI admin (pamoja na data zake zote) — kwa BATCHES
        # (usifute waliobaki: endapo watumiaji ni zaidi ya max_delete).
        from bson import ObjectId as _OID
        batch = max_delete
        total_users = total_events = total_views = 0
        while True:
            cur = db.users.find({"is_admin": {"$ne": True}}, {"_id": 1}).limit(batch)
            ids = [str(u["_id"]) async for u in cur]
            if not ids:
                break
            oids = [_OID(i) for i in ids]
            total_users += (await db.users.delete_many({"_id": {"$in": oids}})).deleted_count
            total_events += (await db.event_log.delete_many({"actor_user_id": {"$in": ids}})).deleted_count
            total_views += (await db.page_views.delete_many({"user_id": {"$in": ids}})).deleted_count
            await db.matches.delete_many({"$or": [{"user_a_id": {"$in": ids}}, {"user_b_id": {"$in": ids}}]})
            await db.notifications.delete_many({"user_id": {"$in": ids}})
            await db.messages.delete_many({"$or": [{"from_user_id": {"$in": ids}}, {"to_user_id": {"$in": ids}}]})
            await db.call_logs.delete_many({"$or": [{"from_user_id": {"$in": ids}}, {"to_user_id": {"$in": ids}}]})
            # Orphan auth records (OTP / resets / verifications) za waliofutwa
            await db.login_otps.delete_many({"user_id": {"$in": oids}})
            await db.password_resets.delete_many({"user_id": {"$in": oids}})
            await db.email_verifications.delete_many({"user_id": {"$in": oids}})
            if len(ids) < batch:
                break
        await _bust_admin_caches()
        return {"ok": True, "mode": "wipe_all",
                "deleted_users": total_users,
                "deleted_events": total_events, "deleted_views": total_views}

    # MAKINI: ni patterns za majaribio ZILIZOJULIKANA tu (jina linaloanza na
    # kitu hiki au lenye neno kamili) — usifute mtu halisi kwa kosa. `\b` pande
    # zote mbili: "Protest" HAIENDANI (test si neno kamili), "Test Debug" ndiyo.
    test_re = re.compile(
        r"\b(test|demo|debug|dummy|sample|zzz)\b|\b(CO|Mwalimu|RN|EN|ANO|NO|MA|ACO|CA|AMO|MD|LAB|PHARM|HA) — ",
        re.I,
    )
    test_phone = re.compile(r"^(\+?255)?0{2,}|^(\+?255)?70000")
    q = {"is_admin": {"$ne": True},
         "$or": [{"full_name": test_re}, {"phone_primary": test_phone}]}
    cur = db.users.find(q, {"full_name": 1, "phone_primary": 1, "created_at": 1}).sort("created_at", 1).limit(max_delete)
    ids = []
    async for u in cur:
        ids.append(str(u["_id"]))
    if not ids:
        return {"ok": True, "deleted_users": 0, "deleted_events": 0, "deleted_views": 0}
    from bson import ObjectId as _OID
    oids = [_OID(i) for i in ids]
    del_users = await db.users.delete_many({"_id": {"$in": oids}})
    del_events = await db.event_log.delete_many({"actor_user_id": {"$in": ids}})
    del_views = await db.page_views.delete_many({"user_id": {"$in": ids}})
    await db.matches.delete_many({"$or": [{"user_a_id": {"$in": ids}}, {"user_b_id": {"$in": ids}}]})
    await db.notifications.delete_many({"user_id": {"$in": ids}})
    await _bust_admin_caches()
    return {"ok": True, "deleted_users": del_users.deleted_count,
            "deleted_events": del_events.deleted_count, "deleted_views": del_views.deleted_count}


class PageViewIn(BaseModel):
    path: str
    referrer: str | None = None


@router.post("/page-view")
async def log_page_view(body: PageViewIn, user=Depends(current_user)):
    """Frontend calls this on route change. Persisted directly (reliable — the
    admin reports read `page_views`) AND published as an event for the audit
    stream / live consumers."""
    now = datetime.now(timezone.utc)
    await get_db().page_views.insert_one({
        "user_id": str(user["_id"]),
        "user_name": user["full_name"],
        "path": body.path,
        "referrer": body.referrer,
        "visited_at": now,
    })
    publish(TOPIC_PAGE_VIEWED, {
        "event": "page.viewed", "user_id": str(user["_id"]),
        "user_name": user["full_name"], "path": body.path,
        "referrer": body.referrer, "occurred_at": now.isoformat(),
    })
    return {"ok": True}


@router.get("/monitoring")
async def monitoring(_=Depends(current_admin)):
    """Snapshot of key backend metrics + broker config (redacted)."""
    db = get_db()
    return {
        "process_started_at": os.environ.get("_KV_STARTED_AT"),
        "mongo": {"uri_host": settings.mongo_uri.split("@")[-1].split("/")[0]},
        "mqtt": {"host": settings.mqtt_host, "port": settings.mqtt_port, "tls": settings.mqtt_use_tls,
                 "username_set": bool(settings.mqtt_username)},
        "redis_url": settings.redis_url,
        "collections": {
            "users": await db.users.count_documents({}),
            "matches": await db.matches.count_documents({}),
            "messages": await db.messages.count_documents({}),
            "call_logs": await db.call_logs.count_documents({}),
            "event_log": await db.event_log.count_documents({}),
        },
        "csv_dir": settings.csv_output_dir,
    }


# ── Password Reset Requests (admin approve/reject) ──────────────────

@router.get("/password-resets")
async def list_password_resets(
    status: str = Query("pending", description="pending|approved|rejected"),
    bypass_cache: bool = Query(False),
    admin=Depends(current_admin),
):
    """List password reset requests filtered by status."""
    cache_key = f"admin:password_resets:{status}"
    if not bypass_cache:
        cached = await _cache_get(cache_key)
        if cached is not None:
            return cached
    db = get_db()
    query = {"status": status} if status else {}
    cursor = db.password_resets.find(query).sort("created_at", -1).limit(100)
    items = []
    async for doc in cursor:
        # Enrich with user info
        user = await db.users.find_one({"_id": doc["user_id"]}, {"full_name": 1, "phone_primary": 1, "email": 1}) or {}
        items.append({
            "id": str(doc["_id"]),
            "user_id": str(doc["user_id"]),
            "full_name": doc.get("full_name") or user.get("full_name", "—"),
            "phone": doc.get("phone", "—"),
            "email": user.get("email", ""),
            "status": doc.get("status", "pending"),
            "created_at": doc.get("created_at", "").isoformat() if hasattr(doc.get("created_at", ""), "isoformat") else str(doc.get("created_at", "")),
        })
    counts = {
        "pending": await db.password_resets.count_documents({"status": "pending"}),
        "approved": await db.password_resets.count_documents({"status": "approved"}),
        "rejected": await db.password_resets.count_documents({"status": "rejected"}),
    }
    result = {"items": items, "counts": counts}
    await _cache_set(cache_key, result, 10)
    return result


@router.post("/password-resets/{reset_id}/approve")
async def approve_password_reset(
    reset_id: str,
    admin=Depends(current_admin),
):
    """Admin approves password reset — marks as approved and logs the code for SMS."""
    db = get_db()
    try:
        oid = ObjectId(reset_id)
    except Exception:
        raise HTTPException(400, "ID batili")
    doc = await db.password_resets.find_one({"_id": oid})
    if not doc:
        raise HTTPException(404, "Ombi halijapatikana")
    if doc.get("status") != "pending":
        raise HTTPException(400, "Ombi tayari limeshachukuliwa")
    await db.password_resets.update_one({"_id": oid}, {"$set": {"status": "approved", "approved_at": datetime.now(timezone.utc), "approved_by": str(admin["_id"])}})
    publish(TOPIC_USER_PASSWORD_RESET_REQUESTED, {
        "event": "user.password_reset_approved",
        "user_id": str(doc["user_id"]),
        "reset_id": reset_id,
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    # REAL-TIME: push WS event to the user so they know immediately
    await _push_ws(str(doc["user_id"]), {
        "event": "user.password_reset_approved",
        "user_id": str(doc["user_id"]),
        "reset_id": reset_id,
        "message": "Ombi lako la kubadilisha password limekubaliwa. Weka password mpya sasa.",
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })    # Bust cache ya password resets — data mpya ionekane papo hapo
    try:
        r = get_redis()
        keys = [k async for k in r.scan_iter("admin:password_resets:*")]
        if keys:
            await r.delete(*keys)
    except Exception:
        pass
    return {"ok": True, "message": "Ombi limekubaliwa"}



@router.post("/password-resets/{reset_id}/reject")
async def reject_password_reset(
    reset_id: str,
    admin=Depends(current_admin),
):
    """Admin rejects password reset request."""
    db = get_db()
    try:
        oid = ObjectId(reset_id)
    except Exception:
        raise HTTPException(400, "ID batili")
    doc = await db.password_resets.find_one({"_id": oid})
    if not doc:
        raise HTTPException(404, "Ombi halijapatikana")
    if doc.get("status") != "pending":
        raise HTTPException(400, "Ombi tayari limeshachukuliwa")
    await db.password_resets.update_one({"_id": oid}, {"$set": {"status": "rejected", "rejected_at": datetime.now(timezone.utc), "rejected_by": str(admin["_id"])}})
    publish(TOPIC_USER_PASSWORD_RESET_REQUESTED, {
        "event": "user.password_reset_rejected",
        "user_id": str(doc["user_id"]),
        "reset_id": reset_id,
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    await _push_ws(str(doc["user_id"]), {
        "event": "user.password_reset_rejected",
        "user_id": str(doc["user_id"]),
        "reset_id": reset_id,
        "message": "Ombi lako la kubadilisha password limekataliwa.",
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })    # Bust cache ya password resets
    try:
        r = get_redis()
        keys = [k async for k in r.scan_iter("admin:password_resets:*")]
        if keys:
            await r.delete(*keys)
    except Exception:
        pass
    return {"ok": True, "message": "Ombi limekataliwa"}


# ─── Contact Permission Settings ───────────────────────────────────────────
# Global: require_payment_for_contact (admin anaweza kuzima kwa wote)
# Per-user: contact_enabled (admin anaweza kumruhusu mtu binafsi)

class ContactSettings(BaseModel):
    require_payment: bool


@router.get("/settings/contact")
async def get_contact_settings(admin=Depends(current_admin)):
    db = get_db()
    doc = await db.settings.find_one({"key": "contact"})
    return {
        "require_payment": bool(doc.get("require_payment", True)) if doc else True,
    }


@router.put("/settings/contact")
async def update_contact_settings(body: ContactSettings, admin=Depends(current_admin)):
    db = get_db()
    await db.settings.update_one(
        {"key": "contact"},
        {"$set": {"require_payment": body.require_payment, "updated_at": datetime.now(timezone.utc)}},
        upsert=True,
    )
    return {"ok": True, "require_payment": body.require_payment}


@router.patch("/users/{user_id}/contact-toggle")
async def toggle_user_contact(user_id: str, admin=Depends(current_admin)):
    db = get_db()
    user = await db.users.find_one({"_id": ObjectId(user_id)})
    if not user:
        raise HTTPException(404, "Mtumiaji hapo")
    new_val = not user.get("contact_enabled", False)
    await db.users.update_one(
        {"_id": ObjectId(user_id)},
        {"$set": {"contact_enabled": new_val, "updated_at": datetime.now(timezone.utc)}},
    )
    # WS push — mtumiaji apate mabadiliko papo hapo
    await _push_ws(user_id, {
        "event": "contact.toggled",
        "contact_enabled": new_val,
        "occurred_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"ok": True, "contact_enabled": new_val}
